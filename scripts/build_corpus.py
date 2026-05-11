"""
build_corpus.py — Generates binary test fixtures for the Musterbau GmbH RAG test corpus.

Source of truth: docs/tests/phase3/corpus/MUSTERBAU.md (frozen v1.0, 2026-05-08).
All numbers, names, and IDs are taken verbatim from that document.

Deps:
    pip install reportlab>=4.0 python-docx>=1.1.0 openpyxl>=3.1.0 xlwt>=1.3.0
    OR (preferred, using the backend venv):
    uv pip install -e backend[corpus]

Run:
    python scripts/build_corpus.py
"""

import os
from pathlib import Path

# ---------------------------------------------------------------------------
# Frozen constants from MUSTERBAU.md — do not invent values
# ---------------------------------------------------------------------------

FIRMA = "Musterbau GmbH"
HRB = "HRB 28471 Dortmund"
GRUENDUNGSJAHR = 2003
SITZ = "Dortmund, Nordrhein-Westfalen"
UMSATZ_2025 = 18_450_000
FTE = 127
STEUER_ID = "DE 287 441 903"
GF = "Werner Kahlert"
CFO = "Dr. Monika Steinhoff"
CFO_START = "01.04.2025"

# Bilanz (Abschnitt 6.1)
BILANZ_AKTIVA = 9_050_000
BILANZ_PASSIVA = 9_050_000
ANLAGEVERMOEGEN = 4_820_000
SACHANLAGEN = 3_950_000
IMMAT_VERMOEGEN = 870_000
UMLAUFVERMOEGEN = 4_230_000
VORRAETE = 940_000
FORDERUNGEN = 2_610_000
KASSENBESTAND = 680_000

EK_GESAMT = 4_230_000
STAMMKAPITAL = 250_000
KAPITALRUECKLAGE = 1_200_000
GEWINNRUECKLAGEN = 1_850_000
JAHRESUEBERSCHUSS = 930_000
FK_GESAMT = 4_820_000
LFR_VERBINDLICHKEITEN = 2_900_000
KFR_VERBINDLICHKEITEN = 1_530_000
RUECKSTELLUNGEN = 390_000

# GuV (Abschnitt 6.2)
UMSATZERLOESE = 18_450_000
SONST_ERTRAGE = 180_000
GESAMTLEISTUNG = 18_630_000
MATERIALAUFWAND = -5_540_000
PERSONALAUFWAND = -7_196_000
ABSCHREIBUNGEN = -480_000
SONST_AUFWENDUNGEN = -3_864_000
EBIT = 1_550_000
ZINSEN = -100_000
EBT = 1_450_000
STEUERN = -520_000

# Kapitalflussrechnung (Abschnitt 6.3)
KFR_ANFANGSBESTAND = 780_000
KFR_ENDBESTAND = 680_000
CASHFLOW_LAUFEND = 1_050_000
CASHFLOW_INVEST = -750_000
CASHFLOW_FINANZ = -400_000
NETTO_CASHFLOW = -100_000

# Schlüsselpersonen
SCHLUESSELPERSONEN = [
    ("MA-001", "Werner Kahlert", "Geschäftsführer"),
    ("MA-002", "Dr. Monika Steinhoff", "CFO"),
    ("MA-006", "Sabine Rühle", "Vertriebsleiterin"),
    ("MA-007", "Dirk Hammerschmidt", "Key-Account Manager"),
    ("MA-011", "Tobias Wernecke", "Engineering-Leiter"),
    ("MA-016", "Franz-Josef Metzler", "Produktionsleiter"),
    ("MA-022", "Inga Hollmann", "Verwaltungsleiterin / HR-Lead"),
    ("MA-023", "Benedikt Falk", "Buchhaltung"),
]

# Großkunden (Abschnitt 4.3)
GROSSKUNDEN = [
    ("KD-007", "Bramkamp Industrietechnik GmbH", "ANL", "NRW", 742_000, "MA-007 Dirk Hammerschmidt"),
    ("KD-034", "Rheinische Anlagenbau AG", "ANL", "NRW", 683_000, "MA-007 Dirk Hammerschmidt"),
    ("KD-021", "Bayerische Stahlbau GmbH", "MAS", "BAY", 561_000, "MA-004 Holger Brandt"),
    ("KD-015", "Stadtwerke Gelsenkirchen AöR", "ENE", "NRW", 437_000, "MA-009 Carsten Voigt"),
]

# BM25-Terme (Abschnitt 8) — assigned target files
TERM_BRANDNER = "Brandner-Rücklage"
TERM_FLAECHENNUTZUNG = "Flächennutzungsindex NRW-3"
TERM_SAUERLANDSTERN = "Sauerlandstern-Beschluss"
TERM_POLYMERINTEGRATION = "Polymerintegration Typ IV"
TERM_OBELISK = "OBELISK-7"

# Corpus output root
CORPUS_ROOT = Path(__file__).parent.parent / "docs" / "tests" / "phase3" / "corpus"

# ---------------------------------------------------------------------------
# Colour palette (XQT5 / Musterbau house colours)
# ---------------------------------------------------------------------------
ORANGE = (238 / 255, 127 / 255, 0 / 255)   # #ee7f00
NAVY = (33 / 255, 52 / 255, 82 / 255)       # #213452


# ===========================================================================
# 1. finanzen_2025.xlsx
# ===========================================================================

def build_finanzen(out_path: Path) -> None:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    wb = Workbook()

    # ---- helpers -----------------------------------------------------------
    HEADER_FILL = PatternFill("solid", fgColor="213452")
    HEADER_FONT = Font(bold=True, color="FFFFFF", name="Calibri")
    BOLD_FONT = Font(bold=True, name="Calibri")
    NORMAL_FONT = Font(name="Calibri")
    TOTAL_FILL = PatternFill("solid", fgColor="EE7F00")
    TOTAL_FONT = Font(bold=True, color="FFFFFF", name="Calibri")

    double_bottom = Border(
        bottom=Side(style="double"),
        top=Side(style="thin"),
    )

    def set_header(ws, headers):
        ws.append(headers)
        for cell in ws[1]:
            cell.font = HEADER_FONT
            cell.fill = HEADER_FILL
            cell.alignment = Alignment(horizontal="center")

    def set_total_row(ws, row_idx):
        for cell in ws[row_idx]:
            cell.font = TOTAL_FONT
            cell.fill = TOTAL_FILL
            cell.border = double_bottom

    def set_col_width(ws, col, width):
        ws.column_dimensions[get_column_letter(col)].width = width

    def euro(val):
        return val  # store as number; format below

    EUR_FMT = '#,##0 "€"'

    # ---- Sheet 1: Bilanz ---------------------------------------------------
    ws1 = wb.active
    ws1.title = "Bilanz"

    # Metadata / context header
    ws1["A1"] = f"{FIRMA} — Bilanz zum 31.12.2025"
    ws1["A1"].font = Font(bold=True, size=13, name="Calibri")
    ws1["A2"] = f"HRB: {HRB} | Steuer-ID: {STEUER_ID}"
    ws1["A2"].font = Font(italic=True, name="Calibri")
    # Note: Brandner-Rücklage was re-evaluated at CFO change E-02
    ws1["A3"] = (
        f"Hinweis: {TERM_BRANDNER} wurde im Zuge des CFO-Wechsels (E-02, {CFO_START}) "
        f"neu bewertet. Verantwortlich: {CFO} (MA-002), {SITZ}."
    )
    ws1["A3"].font = Font(italic=True, color="CC0000", name="Calibri")
    ws1["A4"] = (
        f"Externer Kennwert: {TERM_FLAECHENNUTZUNG} — gemäß Jahresabschlussbesprechung E-08 "
        f"(09.12.2025) protokolliert von MA-022 Inga Hollmann und MA-023 Benedikt Falk."
    )
    ws1["A4"].font = Font(italic=True, name="Calibri")

    # Aktiva section
    aktiva_start = 6
    ws1[f"A{aktiva_start}"] = "AKTIVA"
    ws1[f"A{aktiva_start}"].font = Font(bold=True, size=12, name="Calibri")

    aktiva_rows = [
        ("Position", "Betrag (€)"),
        ("Anlagevermögen gesamt", ANLAGEVERMOEGEN),
        ("  — Sachanlagen (Maschinen, Gebäude)", SACHANLAGEN),
        ("  — Immaterielle Vermögenswerte (PIMS-Lizenzen, Software)", IMMAT_VERMOEGEN),
        ("Umlaufvermögen gesamt", UMLAUFVERMOEGEN),
        ("  — Vorräte", VORRAETE),
        ("  — Forderungen aus Lieferungen und Leistungen", FORDERUNGEN),
        ("  — Kassenbestand und Bankguthaben", KASSENBESTAND),
        ("Bilanzsumme Aktiva", BILANZ_AKTIVA),
    ]

    for i, (pos, val) in enumerate(aktiva_rows):
        r = aktiva_start + 1 + i
        ws1.cell(r, 1, pos)
        ws1.cell(r, 2, val if isinstance(val, int) else val)
        if i == 0:
            ws1.cell(r, 1).font = HEADER_FONT
            ws1.cell(r, 1).fill = HEADER_FILL
            ws1.cell(r, 2).font = HEADER_FONT
            ws1.cell(r, 2).fill = HEADER_FILL
        else:
            ws1.cell(r, 2).number_format = EUR_FMT
            if pos.startswith("Bilanzsumme"):
                ws1.cell(r, 1).font = BOLD_FONT
                ws1.cell(r, 2).font = TOTAL_FONT
                ws1.cell(r, 2).fill = TOTAL_FILL
                ws1.cell(r, 1).border = double_bottom
                ws1.cell(r, 2).border = double_bottom

    passiva_start = aktiva_start + len(aktiva_rows) + 2
    ws1.cell(passiva_start, 1, "PASSIVA").font = Font(bold=True, size=12, name="Calibri")

    passiva_rows = [
        ("Position", "Betrag (€)"),
        ("Eigenkapital gesamt", EK_GESAMT),
        ("  — Stammkapital", STAMMKAPITAL),
        ("  — Kapitalrücklage", KAPITALRUECKLAGE),
        ("  — Gewinnrücklagen (Vorjahre)", GEWINNRUECKLAGEN),
        (f"  — Jahresüberschuss 2025", JAHRESUEBERSCHUSS),
        ("Fremdkapital gesamt", FK_GESAMT),
        ("  — Langfristige Verbindlichkeiten (Bankdarlehen)", LFR_VERBINDLICHKEITEN),
        ("  — Kurzfristige Verbindlichkeiten", KFR_VERBINDLICHKEITEN),
        ("  — Rückstellungen", RUECKSTELLUNGEN),
        ("Bilanzsumme Passiva", BILANZ_PASSIVA),
    ]

    for i, (pos, val) in enumerate(passiva_rows):
        r = passiva_start + 1 + i
        ws1.cell(r, 1, pos)
        ws1.cell(r, 2, val if isinstance(val, int) else val)
        if i == 0:
            ws1.cell(r, 1).font = HEADER_FONT
            ws1.cell(r, 1).fill = HEADER_FILL
            ws1.cell(r, 2).font = HEADER_FONT
            ws1.cell(r, 2).fill = HEADER_FILL
        else:
            ws1.cell(r, 2).number_format = EUR_FMT
            if pos.startswith("Bilanzsumme"):
                ws1.cell(r, 1).font = BOLD_FONT
                ws1.cell(r, 2).font = TOTAL_FONT
                ws1.cell(r, 2).fill = TOTAL_FILL
                ws1.cell(r, 1).border = double_bottom
                ws1.cell(r, 2).border = double_bottom

    set_col_width(ws1, 1, 58)
    set_col_width(ws1, 2, 18)

    # ---- Sheet 2: GuV ------------------------------------------------------
    ws2 = wb.create_sheet("GuV")

    ws2["A1"] = f"{FIRMA} — Gewinn- und Verlustrechnung 2025"
    ws2["A1"].font = Font(bold=True, size=13, name="Calibri")

    guv_rows = [
        ("Position", "Betrag (€)"),
        ("Umsatzerlöse", UMSATZERLOESE),
        ("Sonstige betriebliche Erträge", SONST_ERTRAGE),
        ("Gesamtleistung", GESAMTLEISTUNG),
        ("Materialaufwand", MATERIALAUFWAND),
        ("Personalaufwand", PERSONALAUFWAND),
        ("Abschreibungen", ABSCHREIBUNGEN),
        ("Sonstige betriebliche Aufwendungen", SONST_AUFWENDUNGEN),
        ("Betriebsergebnis (EBIT)", EBIT),
        ("Zinsen und ähnliche Aufwendungen", ZINSEN),
        ("Ergebnis vor Steuern (EBT)", EBT),
        ("Ertragsteuern", STEUERN),
        ("Jahresüberschuss", JAHRESUEBERSCHUSS),
    ]

    SUBTOTAL_ROWS = {"Gesamtleistung", "Betriebsergebnis (EBIT)", "Ergebnis vor Steuern (EBT)", "Jahresüberschuss"}

    for i, (pos, val) in enumerate(guv_rows):
        r = 3 + i
        ws2.cell(r, 1, pos)
        ws2.cell(r, 2, val if isinstance(val, int) else val)
        if i == 0:
            ws2.cell(r, 1).font = HEADER_FONT
            ws2.cell(r, 1).fill = HEADER_FILL
            ws2.cell(r, 2).font = HEADER_FONT
            ws2.cell(r, 2).fill = HEADER_FILL
        else:
            ws2.cell(r, 2).number_format = EUR_FMT
            if pos in SUBTOTAL_ROWS:
                ws2.cell(r, 1).font = BOLD_FONT
                if pos == "Jahresüberschuss":
                    ws2.cell(r, 2).font = TOTAL_FONT
                    ws2.cell(r, 2).fill = TOTAL_FILL
                    ws2.cell(r, 1).border = double_bottom
                    ws2.cell(r, 2).border = double_bottom
                else:
                    ws2.cell(r, 2).font = BOLD_FONT

    # Note on Q3
    note_r = 3 + len(guv_rows) + 1
    ws2.cell(note_r, 1, "Hinweis: Q3-Umsatz 4.920.000 € (Quartalsziel übertroffen, E-06, 15.09.2025).")
    ws2.cell(note_r, 1).font = Font(italic=True, name="Calibri")
    ws2.cell(note_r + 1, 1, f"Verantwortlich: {CFO} (MA-002), MA-016 Franz-Josef Metzler, MA-022 Inga Hollmann.")
    ws2.cell(note_r + 1, 1).font = Font(italic=True, name="Calibri")

    set_col_width(ws2, 1, 48)
    set_col_width(ws2, 2, 18)

    # ---- Sheet 3: Kapitalflussrechnung -------------------------------------
    ws3 = wb.create_sheet("Kapitalflussrechnung")

    ws3["A1"] = f"{FIRMA} — Kapitalflussrechnung 2025 (indirekte Methode)"
    ws3["A1"].font = Font(bold=True, size=13, name="Calibri")

    kfr_rows = [
        ("Position", "Betrag (€)"),
        ("Jahresüberschuss", JAHRESUEBERSCHUSS),
        ("+ Abschreibungen", 480_000),
        ("+/− Veränderung Forderungen", -310_000),
        ("+/− Veränderung Vorräte", -140_000),
        ("+/− Veränderung kurzfristige Verbindlichkeiten", 90_000),
        ("Cashflow aus laufender Geschäftstätigkeit", CASHFLOW_LAUFEND),
        ("Investitionen in Sachanlagen", -620_000),
        ("Investitionen in immaterielle Vermögenswerte", -130_000),
        ("Cashflow aus Investitionstätigkeit", CASHFLOW_INVEST),
        ("Tilgung Bankdarlehen", -280_000),
        ("Auszahlung Gewinnentnahmen", -120_000),
        ("Cashflow aus Finanzierungstätigkeit", CASHFLOW_FINANZ),
        ("Netto-Cashflow", NETTO_CASHFLOW),
        ("Anfangsbestand Zahlungsmittel (01.01.2025)", KFR_ANFANGSBESTAND),
        ("Endbestand Zahlungsmittel (31.12.2025)", KFR_ENDBESTAND),
    ]

    SUBTOTAL_KFR = {
        "Cashflow aus laufender Geschäftstätigkeit",
        "Cashflow aus Investitionstätigkeit",
        "Cashflow aus Finanzierungstätigkeit",
        "Netto-Cashflow",
        "Endbestand Zahlungsmittel (31.12.2025)",
    }

    for i, (pos, val) in enumerate(kfr_rows):
        r = 3 + i
        ws3.cell(r, 1, pos)
        ws3.cell(r, 2, val if isinstance(val, int) else val)
        if i == 0:
            ws3.cell(r, 1).font = HEADER_FONT
            ws3.cell(r, 1).fill = HEADER_FILL
            ws3.cell(r, 2).font = HEADER_FONT
            ws3.cell(r, 2).fill = HEADER_FILL
        else:
            ws3.cell(r, 2).number_format = EUR_FMT
            if pos in SUBTOTAL_KFR:
                ws3.cell(r, 1).font = BOLD_FONT
                if pos == "Endbestand Zahlungsmittel (31.12.2025)":
                    ws3.cell(r, 2).font = TOTAL_FONT
                    ws3.cell(r, 2).fill = TOTAL_FILL
                    ws3.cell(r, 1).border = double_bottom
                    ws3.cell(r, 2).border = double_bottom
                else:
                    ws3.cell(r, 2).font = BOLD_FONT

    note_r = 3 + len(kfr_rows) + 1
    ws3.cell(note_r, 1,
             f"Endbestand {KFR_ENDBESTAND:,} € stimmt mit Bilanzposition 'Kassenbestand und Bankguthaben' überein.".replace(",", "."))
    ws3.cell(note_r, 1).font = Font(italic=True, name="Calibri")
    ws3.cell(note_r + 1, 1,
             f"Jahresabschlussbesprechung E-08 (09.12.2025): MA-001 {GF}, MA-002 {CFO}, MA-022 Inga Hollmann, MA-023 Benedikt Falk.")
    ws3.cell(note_r + 1, 1).font = Font(italic=True, name="Calibri")

    set_col_width(ws3, 1, 52)
    set_col_width(ws3, 2, 18)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(out_path)


# ===========================================================================
# 2. memo_strategieklausur.docx
# ===========================================================================

def build_memo(out_path: Path) -> None:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    # Title
    title = doc.add_paragraph("Musterbau GmbH — Internes Memorandum", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("Strategieklausur 2025 und Sauerlandstern-Beschluss")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].bold = True
    sub.runs[0].font.size = Pt(13)
    sub.runs[0].font.color.rgb = RGBColor(33, 52, 82)

    doc.add_paragraph(
        "Verteilerliste: MA-001 Werner Kahlert, MA-002 Dr. Monika Steinhoff, "
        "MA-006 Sabine Rühle, MA-011 Tobias Wernecke, MA-022 Inga Hollmann"
    ).italic = True

    doc.add_paragraph(
        "Datum: 14.01.2025 (E-01) und 10.07.2025 (E-05) | Ort: Dortmund / Winterberg"
    ).italic = True

    # ---- Hintergrund -------------------------------------------------------
    doc.add_heading("Hintergrund", level=1)

    doc.add_paragraph(
        f"Die Musterbau GmbH ({HRB}) führte am 14. Januar 2025 ihre jährliche "
        f"Strategieklausur am Firmensitz in {SITZ} durch (Ereignis E-01). "
        f"Unter Leitung von Geschäftsführer {GF} (MA-001) wurden die strategischen "
        f"Schwerpunkte für das Geschäftsjahr 2025 festgelegt und das interne "
        f"Programm {TERM_OBELISK} zur Einführung des PIMS-D-Rollout initiiert."
    )

    doc.add_paragraph(
        f"Am 1. April 2025 trat {CFO} (MA-002) die Nachfolge von Klaus-Dieter Frommann "
        f"als CFO an (Ereignis E-02). Im Zuge dieses Wechsels wurde die "
        f"{TERM_BRANDNER} neu bewertet. Inga Hollmann (MA-022) koordinierte den "
        f"organisatorischen Übergang gemeinsam mit Benedikt Falk (MA-023) aus der Buchhaltung."
    )

    doc.add_paragraph(
        f"Am 10. Juli 2025 fand eine außerordentliche Klausur in Winterberg statt, "
        f"aus der der {TERM_SAUERLANDSTERN} hervorging (Ereignis E-05). "
        f"Teilnehmer waren {GF} (MA-001), {CFO} (MA-002), Sabine Rühle (MA-006), "
        f"Tobias Wernecke (MA-011) und Inga Hollmann (MA-022)."
    )

    # ---- Ergebnisse --------------------------------------------------------
    doc.add_heading("Ergebnisse", level=1)

    doc.add_paragraph(
        f"Folgende Maßnahmen wurden im Rahmen beider Klausuren beschlossen. "
        f"Der {TERM_SAUERLANDSTERN} definiert insbesondere die "
        f"{TERM_POLYMERINTEGRATION} als Pflichtoption für das PIMS-E-Gehäuse."
    )

    # 4-row table: Massnahme | Verantwortlich | Termin | Status
    table = doc.add_table(rows=5, cols=4)
    table.style = "Table Grid"

    headers = ["Maßnahme", "Verantwortlich", "Termin", "Status"]
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        # shade header
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "213452")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:val"), "clear")
        tc_pr.append(shd)

    data_rows = [
        (
            f"PIMS-D-Rollout (Programm {TERM_OBELISK})",
            "MA-011 Tobias Wernecke",
            "Q3 2025",
            "In Bearbeitung",
        ),
        (
            f"Umsetzung {TERM_POLYMERINTEGRATION} für PIMS-E-Gehäuse",
            "MA-011 Tobias Wernecke, MA-016 Franz-Josef Metzler",
            "31.10.2025",
            "Freigegeben",
        ),
        (
            f"Neubewertung {TERM_BRANDNER} nach CFO-Wechsel",
            "MA-002 Dr. Monika Steinhoff, MA-023 Benedikt Falk",
            "30.04.2025",
            "Abgeschlossen",
        ),
        (
            "Rahmenvertrag Großauftrag KD-007 (Bramkamp Industrietechnik GmbH)",
            "MA-007 Dirk Hammerschmidt",
            "18.04.2025",
            "Unterzeichnet",
        ),
    ]

    for row_idx, (massnahme, verantw, termin, status) in enumerate(data_rows):
        row = table.rows[row_idx + 1]
        row.cells[0].text = massnahme
        row.cells[1].text = verantw
        row.cells[2].text = termin
        row.cells[3].text = status

    doc.add_paragraph("")  # spacing

    # ---- Nächste Schritte --------------------------------------------------
    doc.add_heading("Nächste Schritte", level=1)

    doc.add_paragraph(
        f"1. Tobias Wernecke (MA-011) legt bis 31.08.2025 einen detaillierten "
        f"Projektplan für {TERM_OBELISK} vor, der die Anforderungen der "
        f"Großkunden KD-007 (Bramkamp Industrietechnik GmbH, 742.000 € p.a.) "
        f"und KD-034 (Rheinische Anlagenbau AG, 683.000 € p.a.) berücksichtigt."
    )

    doc.add_paragraph(
        f"2. Dr. Monika Steinhoff (MA-002) übergibt die aktualisierte Bewertung "
        f"der {TERM_BRANDNER} bis 15.05.2025 an die Geschäftsführung. "
        f"Benedikt Falk (MA-023) bereitet die buchhalterischen Unterlagen vor."
    )

    doc.add_paragraph(
        f"3. Sabine Rühle (MA-006) koordiniert mit Dirk Hammerschmidt (MA-007) "
        f"die Vertriebsstrategie für die Zielkunden KD-021 (Bayerische Stahlbau GmbH, "
        f"561.000 €) und KD-015 (Stadtwerke Gelsenkirchen AöR, 437.000 €) im "
        f"Zusammenhang mit dem {TERM_SAUERLANDSTERN}."
    )

    doc.add_paragraph(
        f"4. Inga Hollmann (MA-022) koordiniert die HR-seitige Begleitung des "
        f"CFO-Übergangs und stellt sicher, dass alle Vollmachten und Zeichnungsrechte "
        f"von Klaus-Dieter Frommann auf {CFO} (MA-002) übertragen werden."
    )

    doc.add_paragraph(
        f"5. Das Engineering-Team (MA-011 Wernecke, MA-012 Sven Unterberg, "
        f"MA-013 Claudia Hölscher) erarbeitet die technische Spezifikation für "
        f"die {TERM_POLYMERINTEGRATION} bis 30.09.2025."
    )

    doc.add_paragraph("")
    foot = doc.add_paragraph(
        f"Dieses Memorandum ist vertraulich. Musterbau GmbH, {SITZ}. "
        f"Dokumentversion: 1.0 — erstellt im Auftrag von {GF} (MA-001)."
    )
    foot.runs[0].italic = True
    foot.runs[0].font.size = Pt(9)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


# ===========================================================================
# Logo drawing helper (shared between geschaeftsbericht + phash fixtures)
# ===========================================================================

def _draw_logo(canvas, x, y, width=120, height=30):
    """
    Draws a deterministic 'M GmbH' logo box at (x, y).
    Same geometry on every call → identical pHash across pages.
    Navy rectangle background, white bold 'M GmbH' text, orange left stripe.
    """
    from reportlab.lib.colors import HexColor
    navy = HexColor("#213452")
    orange = HexColor("#ee7f00")
    white = HexColor("#ffffff")

    canvas.saveState()
    # Navy background
    canvas.setFillColor(navy)
    canvas.rect(x, y, width, height, fill=1, stroke=0)
    # Orange left accent stripe
    canvas.setFillColor(orange)
    canvas.rect(x, y, 8, height, fill=1, stroke=0)
    # White label
    canvas.setFillColor(white)
    canvas.setFont("Helvetica-Bold", 11)
    canvas.drawString(x + 14, y + 10, "M GmbH")
    canvas.restoreState()


def _draw_page_header(canvas, page_width, page_height, title_text):
    """Draws the standardised page header with logo + document title."""
    from reportlab.lib.colors import HexColor
    navy = HexColor("#213452")

    # Logo — fixed position top-right
    _draw_logo(canvas, page_width - 150, page_height - 50)

    # Header line
    canvas.saveState()
    canvas.setStrokeColor(navy)
    canvas.setLineWidth(1.5)
    canvas.line(40, page_height - 60, page_width - 40, page_height - 60)
    # Document title top-left
    canvas.setFillColor(navy)
    canvas.setFont("Helvetica-Bold", 9)
    canvas.drawString(40, page_height - 50, title_text)
    canvas.restoreState()


def _draw_page_footer(canvas, page_width, page_num, total_pages=None):
    from reportlab.lib.colors import HexColor
    grey = HexColor("#555555")
    canvas.saveState()
    canvas.setStrokeColor(HexColor("#213452"))
    canvas.setLineWidth(0.5)
    canvas.line(40, 40, page_width - 40, 40)
    canvas.setFillColor(grey)
    canvas.setFont("Helvetica", 8)
    canvas.drawString(40, 28, f"Musterbau GmbH | HRB 28471 Dortmund | Vertraulich")
    pg_text = f"Seite {page_num}" + (f" von {total_pages}" if total_pages else "")
    canvas.drawRightString(page_width - 40, 28, pg_text)
    canvas.restoreState()


# ===========================================================================
# 3. geschaeftsbericht_2025.pdf
# ===========================================================================

def build_geschaeftsbericht(out_path: Path) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.colors import HexColor, white
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak,
        HRFlowable,
    )
    from reportlab.graphics.shapes import Drawing, Rect, String, Line
    from reportlab.graphics.charts.barcharts import VerticalBarChart
    from reportlab.graphics import renderPDF

    PAGE_W, PAGE_H = A4
    MARGIN = 2.5 * cm
    NAVY_RL = HexColor("#213452")
    ORANGE_RL = HexColor("#ee7f00")

    styles = getSampleStyleSheet()

    def sty(name, **kw):
        s = ParagraphStyle(name, parent=styles["Normal"], **kw)
        return s

    H1 = sty("H1gb", fontSize=16, textColor=NAVY_RL, spaceBefore=18, spaceAfter=6,
              fontName="Helvetica-Bold", leading=20)
    H2 = sty("H2gb", fontSize=12, textColor=NAVY_RL, spaceBefore=12, spaceAfter=4,
              fontName="Helvetica-Bold", leading=15)
    BODY = sty("Bodygb", fontSize=10, leading=14, spaceAfter=8)
    CAPTION = sty("Capgb", fontSize=8, leading=11, textColor=HexColor("#555555"),
                  spaceAfter=12, alignment=1)  # centred

    doc_title = f"{FIRMA} — Geschäftsbericht 2025"

    # We use a canvas-based approach with onFirstPage / onLaterPages
    from reportlab.pdfgen import canvas as rl_canvas
    import io

    out_path.parent.mkdir(parents=True, exist_ok=True)

    def make_doc(path):
        return SimpleDocTemplate(
            str(path),
            pagesize=A4,
            leftMargin=MARGIN, rightMargin=MARGIN,
            topMargin=3.5 * cm, bottomMargin=2.5 * cm,
            title=doc_title,
            author=FIRMA,
        )

    # ---- bar chart (unique figure on Finanzkennzahlen page) ----------------
    def make_bar_chart_drawing():
        d = Drawing(380, 160)
        bc = VerticalBarChart()
        bc.x = 40
        bc.y = 20
        bc.width = 300
        bc.height = 120
        bc.data = [[EBIT / 1000, JAHRESUEBERSCHUSS / 1000, EBT / 1000]]
        bc.categoryAxis.categoryNames = ["EBIT", "Jahresüberschuss", "EBT"]
        bc.bars[0].fillColor = ORANGE_RL
        bc.valueAxis.valueMin = 0
        bc.valueAxis.valueMax = 2000
        bc.valueAxis.valueStep = 500
        bc.valueAxis.labelTextFormat = "%d T€"
        d.add(bc)
        lbl = String(190, 150, "Ergebnisübersicht 2025 (in T€)", textAnchor="middle",
                     fontSize=9, fillColor=NAVY_RL, fontName="Helvetica-Bold")
        d.add(lbl)
        return d

    # ---- cover page --------------------------------------------------------
    story = []

    # Spacer for header clearance
    story.append(Spacer(1, 1.5 * cm))

    # Cover title block
    cover_data = [
        [Paragraph(f"<font color='#213452'><b>Geschäftsbericht 2025</b></font>",
                   ParagraphStyle("cov", fontSize=28, leading=34, fontName="Helvetica-Bold"))],
        [Paragraph(FIRMA, ParagraphStyle("covf", fontSize=18, leading=22,
                                         textColor=ORANGE_RL, fontName="Helvetica-Bold",
                                         spaceAfter=6))],
        [Paragraph(f"{HRB} | {SITZ}", ParagraphStyle("covs", fontSize=10,
                                                       textColor=HexColor("#555555")))],
    ]
    ct = Table(cover_data, colWidths=[PAGE_W - 2 * MARGIN])
    ct.setStyle(TableStyle([
        ("ALIGN", (0, 0), (-1, -1), "LEFT"),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(ct)
    story.append(Spacer(1, 0.8 * cm))
    story.append(HRFlowable(width="100%", thickness=2, color=ORANGE_RL))
    story.append(Spacer(1, 0.5 * cm))

    story.append(Paragraph(
        f"Geschäftsführer: {GF} (MA-001) | CFO: {CFO} (MA-002, seit {CFO_START})",
        BODY))
    story.append(Paragraph(
        f"Mitarbeitende: {FTE} FTE | Umsatz 2025: {UMSATZ_2025:,} €".replace(",", "."),
        BODY))
    story.append(Paragraph(
        f"Steuer-ID: {STEUER_ID} | Gründungsjahr: {GRUENDUNGSJAHR}",
        BODY))

    story.append(PageBreak())

    # ---- Vorwort -----------------------------------------------------------
    story.append(Paragraph("1. Vorwort des Geschäftsführers", H1))
    story.append(Paragraph(
        f"Sehr geehrte Geschäftspartnerinnen und Geschäftspartner,", BODY))
    story.append(Paragraph(
        f"das Geschäftsjahr 2025 war für die {FIRMA} ein Jahr des Aufbruchs und der "
        f"strategischen Neuausrichtung. Mit einem Jahresumsatz von "
        f"{UMSATZ_2025:,} € und einem Jahresüberschuss von {JAHRESUEBERSCHUSS:,} € "
        f"haben wir unsere Ziele übertroffen und die Grundlage für weiteres Wachstum "
        f"gelegt.".replace(",", "."),
        BODY))
    story.append(Paragraph(
        f"Besondere Meilensteine des Jahres waren die Unterzeichnung des Rahmenvertrags "
        f"mit der Bramkamp Industrietechnik GmbH (KD-007) über 742.000 € p.a. sowie "
        f"die Rezertifizierung nach ISO 9001. Das strategische Programm {TERM_OBELISK} "
        f"zur Einführung von PIMS-D wurde planmäßig initiiert.",
        BODY))
    story.append(Paragraph(
        f"Ich danke allen {FTE} Mitarbeitenden für ihren außerordentlichen Einsatz.",
        BODY))
    story.append(Paragraph(
        f"Werner Kahlert (MA-001), Geschäftsführer | {SITZ}, Dezember 2025",
        sty("sig", fontSize=9, textColor=HexColor("#555555"))))

    story.append(PageBreak())

    # ---- Lagebericht -------------------------------------------------------
    story.append(Paragraph("2. Lagebericht", H1))

    story.append(Paragraph("2.1 Gesamtwirtschaftliches Umfeld", H2))
    story.append(Paragraph(
        "Das Geschäftsjahr 2025 verlief für den deutschen Maschinen- und Anlagenbau "
        "insgesamt stabil. Die Nachfrage nach integrierten Überwachungslösungen — "
        "insbesondere im Bereich Predictive Maintenance — stieg weiter an. "
        "Musterbau GmbH konnte von dieser Entwicklung profitieren und ihre Position "
        "als Anbieter von PIMS-Lösungen (Prozessintegriertes Monitoring-System) weiter festigen.",
        BODY))

    story.append(Paragraph("2.2 Schlüsselereignisse", H2))
    story.append(Paragraph(
        f"<b>E-01 — Strategieklausur (14.01.2025):</b> Unter Leitung von {GF} "
        f"(MA-001) wurden die Jahresziele festgelegt und das Programm {TERM_OBELISK} "
        f"(PIMS-D-Rollout) beschlossen. Teilnehmer: MA-001, MA-006 Sabine Rühle, "
        f"MA-011 Tobias Wernecke, MA-016 Franz-Josef Metzler, MA-022 Inga Hollmann.",
        BODY))
    story.append(Paragraph(
        f"<b>E-02 — CFO-Wechsel (01.04.2025):</b> {CFO} (MA-002) übernahm die "
        f"CFO-Funktion von Klaus-Dieter Frommann. Inga Hollmann (MA-022) begleitete "
        f"den Übergang organisatorisch.",
        BODY))
    story.append(Paragraph(
        f"<b>E-03 — Großauftrag KD-007 (18.04.2025):</b> Rahmenvertrag PIMS-A + PIMS-D "
        f"mit Bramkamp Industrietechnik GmbH, Laufzeit 3 Jahre, Wert 742.000 € p.a. "
        f"Key-Account: Dirk Hammerschmidt (MA-007).",
        BODY))
    story.append(Paragraph(
        f"<b>E-04 — ISO-9001-Rezertifizierung (22.05.2025):</b> Audit erfolgreich bestanden "
        f"unter Leitung von MA-016 Franz-Josef Metzler und MA-019 Sandra Kemper.",
        BODY))
    story.append(Paragraph(
        f"<b>E-06 — Q3-Review (15.09.2025):</b> Q3-Umsatz 4.920.000 €, Quartalsziel "
        f"übertroffen. Teilnehmer: MA-001, MA-002, MA-006, MA-011, MA-016, MA-022.",
        BODY))
    story.append(Paragraph(
        f"<b>E-08 — Jahresabschlussbesprechung (09.12.2025):</b> Jahresüberschuss "
        f"{JAHRESUEBERSCHUSS:,} € festgestellt. Gewinnverwendungsvorschlag beschlossen.".replace(",", "."),
        BODY))

    story.append(PageBreak())

    # ---- Finanzkennzahlen --------------------------------------------------
    story.append(Paragraph("3. Finanzkennzahlen", H1))

    fin_table_data = [
        ["Kennzahl", "Wert 2025"],
        ["Umsatzerlöse", f"{UMSATZERLOESE:,} €".replace(",", ".")],
        ["Gesamtleistung", f"{GESAMTLEISTUNG:,} €".replace(",", ".")],
        ["EBIT", f"{EBIT:,} €".replace(",", ".")],
        ["EBT", f"{EBT:,} €".replace(",", ".")],
        ["Jahresüberschuss", f"{JAHRESUEBERSCHUSS:,} €".replace(",", ".")],
        ["Bilanzsumme", f"{BILANZ_AKTIVA:,} €".replace(",", ".")],
        ["Eigenkapital", f"{EK_GESAMT:,} €".replace(",", ".")],
        ["Endbestand Zahlungsmittel", f"{KFR_ENDBESTAND:,} €".replace(",", ".")],
        ["Personalaufwand / Umsatz", "39,0 %"],
    ]

    ft = Table(fin_table_data, colWidths=[9 * cm, 7 * cm])
    ft.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), NAVY_RL),
        ("TEXTCOLOR", (0, 0), (-1, 0), white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [HexColor("#f5f5f5"), white]),
        ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#cccccc")),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("FONTNAME", (0, -1), (-1, -1), "Helvetica-Bold"),
    ]))
    story.append(ft)
    story.append(Spacer(1, 0.4 * cm))

    # Unique bar chart figure
    chart_d = make_bar_chart_drawing()
    story.append(chart_d)
    story.append(Paragraph("Abbildung 1: Ergebnisübersicht 2025 (EBIT / Jahresüberschuss / EBT in T€)", CAPTION))

    story.append(PageBreak())

    # ---- Mitarbeitende -----------------------------------------------------
    story.append(Paragraph("4. Mitarbeitende", H1))
    story.append(Paragraph(
        f"Die {FIRMA} beschäftigt zum 31.12.2025 insgesamt {FTE} Vollzeitäquivalente (FTE) "
        f"in fünf Abteilungen am Standort {SITZ}.",
        BODY))

    ma_table_data = [
        ["Abteilung", "Leiter / Leiterin", "FTE"],
        ["Geschäftsführung (GF)", "Werner Kahlert (MA-001)", "3"],
        ["Vertrieb (VT)", "Sabine Rühle (MA-006)", "22"],
        ["Engineering (ENG)", "Tobias Wernecke (MA-011)", "38"],
        ["Produktion (PRO)", "Franz-Josef Metzler (MA-016)", "47"],
        ["Verwaltung (VW)", "Inga Hollmann (MA-022)", "17"],
        ["Gesamt", "", "127"],
    ]
    mat = Table(ma_table_data, colWidths=[7 * cm, 7 * cm, 2.5 * cm])
    mat.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), NAVY_RL),
        ("TEXTCOLOR", (0, 0), (-1, 0), white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTNAME", (0, -1), (-1, -1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("ROWBACKGROUNDS", (0, 1), (-1, -2), [HexColor("#f5f5f5"), white]),
        ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#cccccc")),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
    ]))
    story.append(mat)
    story.append(Spacer(1, 0.3 * cm))
    story.append(Paragraph(
        f"CFO Dr. Monika Steinhoff (MA-002) trat am {CFO_START} die Nachfolge von "
        f"Klaus-Dieter Frommann an (E-02). Die Abteilung Engineering unter Tobias "
        f"Wernecke (MA-011) verantwortet die PIMS-Produktlinie, die mit Programm "
        f"{TERM_OBELISK} eine strategische Erweiterung erfährt.",
        BODY))

    story.append(PageBreak())

    # ---- Ausblick ----------------------------------------------------------
    story.append(Paragraph("5. Ausblick", H1))
    story.append(Paragraph(
        f"Für das Geschäftsjahr 2026 plant die {FIRMA} eine Fortsetzung des "
        f"Wachstumskurses. Schwerpunkte sind die Skalierung von PIMS-D im Rahmen "
        f"des Programms {TERM_OBELISK} sowie der Ausbau des Großkundenportfolios. "
        f"Die vier Großkunden KD-007 (742.000 €), KD-034 (683.000 €), "
        f"KD-021 (561.000 €) und KD-015 (437.000 €) bilden die Basis für ein "
        f"stabiles Umsatzfundament.",
        BODY))
    story.append(Paragraph(
        f"Die strategische Neuausrichtung von PIMS-E, beschlossen im "
        f"{TERM_SAUERLANDSTERN} (E-05), soll im ersten Quartal 2026 in ein "
        f"konkretes Produktangebot überführt werden.",
        BODY))
    story.append(Paragraph(
        f"Verantwortliche Führungskräfte für 2026: {GF} (MA-001), {CFO} (MA-002), "
        f"Sabine Rühle (MA-006), Tobias Wernecke (MA-011), Franz-Josef Metzler (MA-016), "
        f"Inga Hollmann (MA-022).",
        BODY))

    # Build PDF with page headers/footers via a custom canvas
    class HeaderFooterCanvas(rl_canvas.Canvas):
        def __init__(self, filename, **kwargs):
            super().__init__(filename, **kwargs)
            self._page_num = 0

        def showPage(self):
            self._page_num += 1
            self._draw_hf()
            super().showPage()

        def save(self):
            self._page_num += 1
            self._draw_hf()
            super().save()

        def _draw_hf(self):
            _draw_page_header(self, PAGE_W, PAGE_H, doc_title)
            _draw_page_footer(self, PAGE_W, self._page_num)

    d = make_doc(out_path)
    d.build(story, canvasmaker=HeaderFooterCanvas)


# ===========================================================================
# 4. long/handbuch_lang.pdf  (≥ 50 pages)
# ===========================================================================

def build_long_handbook(out_path: Path) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.colors import HexColor, white
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.pdfgen import canvas as rl_canvas

    PAGE_W, PAGE_H = A4
    MARGIN = 2.5 * cm
    NAVY_RL = HexColor("#213452")
    ORANGE_RL = HexColor("#ee7f00")

    styles = getSampleStyleSheet()

    def sty(name, **kw):
        return ParagraphStyle(name, parent=styles["Normal"], **kw)

    H1 = sty("H1hb", fontSize=15, textColor=NAVY_RL, spaceBefore=20, spaceAfter=8,
             fontName="Helvetica-Bold", leading=19)
    H2 = sty("H2hb", fontSize=12, textColor=NAVY_RL, spaceBefore=12, spaceAfter=4,
             fontName="Helvetica-Bold", leading=15)
    H3 = sty("H3hb", fontSize=10, textColor=ORANGE_RL, spaceBefore=8, spaceAfter=3,
             fontName="Helvetica-Bold", leading=13)
    BODY = sty("Bodyhb", fontSize=10, leading=14, spaceAfter=7)
    NOTE = sty("Notehb", fontSize=9, leading=12, spaceAfter=5,
               textColor=HexColor("#444444"), leftIndent=12)

    doc_title = f"{FIRMA} — Technisches Betriebshandbuch PIMS-A (intern)"

    chapters = [
        ("Einleitung und Geltungsbereich", [
            ("Zweck des Dokuments",
             f"Dieses Betriebshandbuch beschreibt den Betrieb, die Wartung und die "
             f"Konfiguration des PIMS-A Basismoduls der {FIRMA} ({HRB}). "
             f"Es richtet sich an autorisiertes Fachpersonal des Engineering-Bereichs "
             f"sowie an qualifizierte Kundentechniker. Verantwortlich: "
             f"Tobias Wernecke (MA-011), Engineering-Leiter."),
            ("Geltungsbereich",
             f"Das Handbuch gilt für alle installierten PIMS-A-Systeme ab Firmware-Version "
             f"3.2.0. Es ergänzt die Kurzanleitung und die technische Spezifikation. "
             f"Produktpreis Listenpreis netto: 48.900 €."),
            ("Normen und Vorschriften",
             "PIMS-A erfüllt folgende Normen: IEC 61508 (Funktionale Sicherheit), "
             "ISO 9001:2015 (Qualitätsmanagementsystem, zertifiziert seit 22.05.2025, E-04), "
             "EN 61000-6-2 (EMV-Immunität Industrie), EN 61000-6-4 (EMV-Emission)."),
            ("Verwendete Symbole",
             "WARNUNG: Sicherheitshinweis, bei Nichtbeachtung Personenschaden möglich. "
             "ACHTUNG: Sachschadenrisiko bei Nichtbeachtung. HINWEIS: Ergänzende Information."),
            ("Dokumentenhistorie",
             "Version 1.0 (2023-01): Erstausgabe. Version 2.0 (2024-06): Erweiterung 24-Kanal. "
             "Version 3.0 (2025-01): Aufnahme Kavitationsschutzprotokoll und Drewermann-Verfahren. "
             "Version 3.1 (2025-09): Aktualisierung nach ISO-Rezertifizierung."),
        ]),
        ("Systemarchitektur", [
            ("Überblick",
             "PIMS-A besteht aus drei Hauptkomponenten: (1) dem Sensorinterface-Modul (SIM) "
             "mit 24 analogen Eingangskanälen, (2) der zentralen Verarbeitungseinheit (CVE) "
             "mit Echtzeit-Betriebssystem, und (3) dem Kommunikationsmodul (KOM) für "
             "Feldbus- und Netzwerkanbindung."),
            ("Sensorinterface-Modul",
             "Das SIM nimmt 24 analoge Signale im Bereich 4–20 mA oder 0–10 V entgegen. "
             "Abtastrate: 1 kHz pro Kanal. Auflösung: 16 Bit. Galvanische Trennung: 500 V. "
             "Schutzart: IP67. Betriebstemperatur: −20 °C bis +60 °C."),
            ("Zentrale Verarbeitungseinheit",
             "Die CVE basiert auf einem ARM Cortex-A53 Quadcore-Prozessor mit 1,4 GHz. "
             "RAM: 2 GB DDR4. Flash: 32 GB eMMC. Betriebssystem: Linux RTOS 5.15 LTS. "
             "Watchdog-Timer: Hardware-basiert, Timeout 500 ms."),
            ("Kommunikationsmodul",
             "Schnittstellen: 2× Ethernet 1 GbE (IEEE 802.3), 1× RS-485 (Modbus RTU), "
             "1× CAN Bus 2.0B, 1× USB 3.0 (Service). Optional: Profibus-DP-Adapter (PIMS-C)."),
            ("Energieversorgung",
             "Nennspannung: 24 VDC (±20 %). Leistungsaufnahme max.: 35 W. "
             "USV-Pufferung: 500 ms Überbrückung bei Versorgungsausfall. "
             "Schutz: Überspannungsschutz Typ 2 nach IEC 61643-11."),
        ]),
        ("Installation und Inbetriebnahme", [
            ("Lieferumfang",
             "1× PIMS-A Hauptgerät im 19\"-Gehäuse (2U), 1× Netzteil 24 VDC/5 A, "
             "1× Ethernet-Patchkabel 2 m (Cat. 6A), 1× USB-Servicekabel, "
             "1× Schnellstartanleitung (DE/EN), 1× Datenträger mit Firmware und Tools."),
            ("Mechanische Montage",
             "Montageart: 19\"-Einschub in Schaltschrank oder Wandmontage mit optionalem "
             "Wandhalter. Mindestabstand zu Wärmequellen: 50 mm oben und unten. "
             "Zulässige Einbaulagen: horizontal ±10°, vertikal ±10°."),
            ("Elektrischer Anschluss",
             "WARNUNG: Vor Anschlussarbeiten Anlage allpolig spannungsfrei schalten. "
             "Klemmenblock X1: 24 VDC +/−. Klemmenblock X2–X25: Analogeingänge Kanal 1–24. "
             "Klemmenblock X26: Digitale E/A (8× DI, 4× DO, 24 VDC)."),
            ("Erstinbetriebnahme",
             "1. Netzspannung einschalten. 2. Bootsequenz abwarten (ca. 45 s). "
             "3. Webinterface aufrufen: http://[IP-Adresse]/setup. 4. Admin-Passwort vergeben. "
             "5. Netzwerkparameter konfigurieren. 6. Kanalparameter eingeben. "
             "7. Funktionstest durchführen (siehe Kapitel 5)."),
            ("Firmware-Update",
             "Updates werden über das Webinterface oder via USB eingespielt. "
             "Voraussetzung: gültige Signatur (SHA-256, Herstellerzertifikat). "
             "Rollback auf vorherige Version innerhalb von 30 Tagen möglich."),
        ]),
        ("Konfiguration", [
            ("Kanalparametrierung",
             "Jeder der 24 Kanäle wird individuell parametriert: Messbereich (z.B. 0–100 bar), "
             "Alarmgrenzen (Warnung + Alarm), Filterzeit (0–10 s), Kanalname (max. 32 Zeichen), "
             "Einheit (frei wählbar, max. 8 Zeichen), Skalierungsformel (linear/quadratisch)."),
            ("Alarmmanagement",
             "PIMS-A unterscheidet vier Alarmstufen: INFO (Stufe 0), WARNUNG (Stufe 1), "
             "ALARM (Stufe 2), KRITISCH (Stufe 3). Jede Stufe hat eigene Relaisausgänge und "
             "Meldekanäle (E-Mail, SNMP-Trap, OPC-UA-Event)."),
            ("Benutzer- und Rollenverwaltung",
             "Rollen: Administrator, Supervisor, Operator, Viewer. "
             "Passwort-Policy: min. 12 Zeichen, Komplexitätsregel, Ablauf 90 Tage. "
             "LDAP/AD-Anbindung optional."),
            ("Datenaufzeichnung",
             "Interner Ringpuffer: 72 Stunden bei 1-s-Auflösung für alle 24 Kanäle. "
             "Langzeitarchiv: NFS/SMB-Share oder S3-Bucket (PIMS-E erforderlich). "
             "Exportformate: CSV, JSON, Parquet."),
            ("Zeitzone und NTP",
             "Zeitzone: UTC+1 (CET) / UTC+2 (CEST) empfohlen. "
             "NTP-Server: bis zu 4 konfigurierbar. Maximale Abweichung für Zeitstempel: ±1 ms."),
        ]),
        ("Prüfung und Kalibrierung", [
            ("Funktionstest nach Inbetriebnahme",
             "Prüfschritte: (1) Selbsttest aller 24 Eingangskanäle mit Kurzschlusstest, "
             "(2) Alarmtest Stufe 1–3, (3) Netzwerktest (Ping, SNMP-Trap-Test), "
             "(4) Datenbanktest (Schreib-/Lesezyklus), (5) Failover-Test (USV-Simulation)."),
            ("Drewermann-Verfahren (Kalibrierung 24 Kanal)",
             "Das Drewermann-Verfahren ist die zugelassene Methode zur Kalibrierung aller "
             "24 Eingangskanäle gemäß E-04 (ISO-9001-Audit, 22.05.2025). "
             "Es sieht eine sequenzielle Beaufschlagung jedes Kanals mit bekannten Referenzgrößen "
             "vor, dokumentiert im Kalibrier-Zertifikat."),
            ("Fensterprotokoll 24-Kanal",
             "Das Fensterprotokoll 24-Kanal definiert die Messfenster für simultane Erfassung. "
             "Fensterbreite: 100 ms (100 Samples je Kanal). Überlappung: 50 ms. "
             "Anwendung: Schallanalyse, Vibrationsüberwachung, Synchronmessung."),
            ("Wiederkehrende Prüfungen",
             "Jährlich: Kalibrierung aller 24 Kanäle nach Drewermann-Verfahren. "
             "Halbjährlich: Alarmtest, USV-Test, Netzwerksicherheitsaudit. "
             "Wöchentlich: Automatischer Selbsttest (firmware-gesteuert, 03:00 Uhr)."),
            ("Prüfprotokoll",
             "Prüfprotokolle sind 10 Jahre aufzubewahren. Format: PDF/A-2b, "
             "signiert mit qualifizierter elektronischer Signatur (QES). "
             "Zuständig: MA-019 Sandra Kemper (Qualitätssicherung)."),
        ]),
        ("Betrieb und Überwachung", [
            ("Normalbetrieb",
             "Im Normalbetrieb erfasst PIMS-A kontinuierlich alle 24 Kanäle, "
             "speichert Messwerte im internen Ringpuffer und stellt sie über das "
             "Webinterface sowie OPC-UA in Echtzeit bereit. Die Status-LED zeigt: "
             "GRÜN = Betrieb OK, GELB = Warnung aktiv, ROT = Alarm aktiv, BLINKEND ROT = Fehler."),
            ("Webinterface",
             "Das Webinterface ist über HTTPS (Port 443) erreichbar. "
             "Funktionen: Echtzeitanzeige aller 24 Kanäle, Trenddiagramme (1 min bis 30 Tage), "
             "Alarmübersicht, Systemstatus, Konfiguration, Firmware-Update."),
            ("Datenexport",
             "Manueller Export: Webinterface → Daten → Export. "
             "Automatischer Export: Cron-Job, konfigurierbar in Webinterface → System → Aufgaben. "
             "Unterstützte Ziele: lokales Laufwerk, NFS, SMB, S3."),
            ("Fernwartung",
             "SSH (Port 22, schlüsselbasiert, Passwort-Auth deaktiviert). "
             "VPN-Empfehlung: IPSec oder WireGuard. "
             "Fernwartungszugang nur für MA-025 Oliver Großmann (IT-Administration) freigegeben."),
            ("Betriebstagebuch",
             "Ereignisse werden automatisch in einem Betriebstagebuch erfasst: "
             "Alarmierungen, Konfigurationsänderungen, Anmeldeereignisse, Firmware-Updates. "
             "Aufbewahrung: 2 Jahre im Gerät, Archiv unbegrenzt auf externem Speicher."),
        ]),
        ("Kavitationsschutzprotokoll", [
            ("Hintergrund",
             "Kavitation in Pumpenkomponenten kann zu Schäden an PIMS-A-Sensoren führen, "
             "die in hydraulischen Anlagen eingesetzt werden. "
             "Das Kavitationsschutzprotokoll wurde im Rahmen des "
             "Produktionsoptimierungs-Workshops (E-07, 03.11.2025) für Linie 3 freigegeben. "
             "Verantwortlich: MA-016 Franz-Josef Metzler, MA-017 Ursula Grentrup."),
            ("Erkennung",
             "Kavitation äußert sich durch hochfrequente Druckschwankungen im Bereich 1–5 kHz. "
             "PIMS-A detektiert diese mit dem Fensterprotokoll 24-Kanal (FFT-Analyse). "
             "Schwellenwert: Pegel > 15 dB über Grundrauschen über 500 ms → Alarm Stufe 2."),
            ("Schutzmaßnahmen",
             "Bei Kavitationsalarm: (1) Automatische Drehzahlreduktion der Pumpe um 15 %, "
             "(2) Protokollierung im Betriebstagebuch, "
             "(3) Benachrichtigung MA-029 Gregor Stenzel (Instandhaltung) per E-Mail. "
             "Dauerhafte Abhilfe: Pumpendruckverhältnis prüfen, Saugseite optimieren."),
            ("Protokollpflicht",
             "Jedes Kavitationsereignis ist in FORM-KSP-2025 zu dokumentieren. "
             "Schwellenwertüberschreitungen > 10 Ereignisse pro Monat: Meldepflicht an MA-016. "
             "Jährliche Auswertung durch MA-019 Sandra Kemper (Qualitätssicherung)."),
            ("Retrofit bestehender Anlagen",
             "Für Bestandsanlagen ohne Kavitationsschutzprotokoll ist ein Firmware-Update "
             "auf mindestens Version 3.0 erforderlich. Kontakt: Tobias Wernecke (MA-011)."),
        ]),
        ("Fehlerdiagnose", [
            ("Fehlercodes E001–E099",
             "E001: Watchdog-Reset. E002: RAM-Fehler (ECC korrigiert). E003: Flash-Fehler. "
             "E010: Kanal-Kurzschluss (Kanal-Nr. im Ereignislog). E011: Kanalunterbrechung. "
             "E020: Netzwerkausfall primäre Schnittstelle. E021: NTP-Synchronisation verloren."),
            ("Fehlercodes E100–E199",
             "E100: Alarmstau (> 100 unquittierte Alarme). E110: Datenbankfehler. "
             "E120: Speicherfüllgrad > 90 %. E130: Zertifikat abgelaufen. "
             "E140: Firmware-Signaturprüfung fehlgeschlagen. E150: USV-Batterie schwach."),
            ("Diagnosewerkzeuge",
             "Systeminformationen: Webinterface → Diagnose → Systembericht. "
             "Kanal-Diagnose: Webinterface → Kanäle → [Kanalname] → Diagnose. "
             "PIMS-Diagnosetool (Windows): PIMS-DT.exe, verfügbar auf Datenträger."),
            ("Häufige Fehlerszenarien",
             "Szenario 1: Keine Daten auf Webinterface → Netzwerk prüfen, Cache leeren. "
             "Szenario 2: Kanal zeigt konstant 0 → Kabelbruch prüfen, Kurzschlusstest. "
             "Szenario 3: Häufige Watchdog-Resets → Betriebstemperatur prüfen, RAM-Test."),
            ("Support und Eskalation",
             f"Erste Anlaufstelle: MA-025 Oliver Großmann (IT-Administration). "
             f"Technische Eskalation: MA-011 Tobias Wernecke (Engineering-Leiter). "
             f"Hersteller-Support: {FIRMA}, {SITZ}, Steuer-ID {STEUER_ID}."),
        ]),
        ("Wartung", [
            ("Wartungsintervalle",
             "Monatlich: Sichtprüfung Kabelverbindungen, Reinigung Lüftungsschlitze. "
             "Vierteljährlich: Prüfung Schraubverbindungen, Prüfung Erdungsanschlüsse. "
             "Jährlich: vollständige Inspektion nach Wartungsplan FORM-WP-2025-A."),
            ("Ersatzteile",
             "Ersatzteile sind ausschließlich vom Hersteller zu beziehen. "
             "Bestellnummern im Ersatzteilkatalog (Dok.-Nr. MB-PIMS-A-EK-2025). "
             "Lagerempfehlung: 1× Netzteilkarte, 1× SIM-Platine, 2× Sicherungen 6,3 A."),
            ("Reinigung",
             "Nur mit trockenen, antistatischen Mitteln reinigen. "
             "Keine Lösungsmittel, keine Druckluft auf Platinen. "
             "Schutzart IP67 gilt nur bei ordnungsgemäß montierten Steckverbindern."),
            ("Außerbetriebnahme",
             "Abmeldung im Kundensystem. Datenlöschung nach BSI-Grundschutz (3-faches Überschreiben). "
             "Entsorgung gemäß WEEE-Richtlinie 2012/19/EU. "
             "Rückmeldung an MA-020 Ralf Overbeck (Lagerhaltung & Logistik)."),
            ("Kalibrierhistorie",
             "Kalibrierungen sind in der PIMS-A-Datenbank und in FORM-KAL-2025 zu dokumentieren. "
             "Exportierbar als PDF/A-2b. Zuständig: MA-019 Sandra Kemper."),
        ]),
        ("Integration mit PIMS-B bis PIMS-E", [
            ("PIMS-B Analytics",
             "PIMS-B (Listenpreis 18.500 € netto) erweitert PIMS-A um statistische Auswertung. "
             "Voraussetzung: PIMS-A Firmware ≥ 3.0. Integration über interne REST-API. "
             "Aktivierung: Lizenzschlüssel im Webinterface eingeben."),
            ("PIMS-C Feldbus-Integration",
             "PIMS-C (Listenpreis 12.200 € netto) integriert Profibus-DP, Modbus-TCP und OPC-UA. "
             "Montage: externer Adapter am Kommunikationsmodul. Konfiguration: Webinterface → Feldbus."),
            ("PIMS-D Predictive Maintenance",
             f"PIMS-D (Listenpreis 29.700 € netto) nutzt historische Sensordaten für KI-gestützte "
             f"Ausfallprognose. Voraussetzung: PIMS-A + PIMS-B. Programm {TERM_OBELISK} (E-01) "
             f"koordiniert den Rollout; Großauftrag KD-007 (742.000 €, E-03) ist Pilotprojekt."),
            ("PIMS-E Cloud-Gateway",
             "PIMS-E (Listenpreis 9.800 € netto) stellt MQTT-basierte Cloud-Anbindung bereit. "
             "Unterstützte Backends: AWS IoT Core, Azure IoT Hub, on-premise MQTT-Broker. "
             "TLS 1.3 obligatorisch. Zertifikatsverwaltung: Let's Encrypt oder eigene PKI."),
            ("Lizenzmanagement",
             "Alle PIMS-Module sind lizenzgebunden. Lizenzen sind an die Seriennummer des "
             "PIMS-A-Hauptgeräts gebunden. Übertragung auf andere Geräte nicht zulässig. "
             "Verwaltung über das Kundenportal oder direkt bei MA-011 Tobias Wernecke."),
        ]),
        ("Sicherheit und Datenschutz", [
            ("IT-Sicherheit",
             "PIMS-A ist nach IEC 62443-3-3 Security Level 2 konzipiert. "
             "Empfehlungen: Netzwerksegmentierung (eigenes VLAN), regelmäßige Firmware-Updates, "
             "Passwort-Rotation, Deaktivierung nicht benötigter Dienste."),
            ("Zugriffskontrolle",
             "Vier Rollen (Administrator, Supervisor, Operator, Viewer) mit differenzierten "
             "Berechtigungen. Jede Anmeldung wird im Betriebstagebuch protokolliert. "
             "Fehlgeschlagene Anmeldeversuche: nach 5 Versuchen 15-Minuten-Sperre."),
            ("Datenschutz (DSGVO)",
             "Personenbezogene Daten (Benutzernamen, Zugriffsprotokolle) werden gemäß DSGVO "
             "verarbeitet. Aufbewahrungsfrist: 2 Jahre. Auskunftsrecht: MA-022 Inga Hollmann."),
            ("Backup und Recovery",
             "Tägliches automatisches Backup der Konfiguration auf externes Ziel. "
             "RPO: 24 Stunden. RTO: 4 Stunden (bei Hardwareersatz). "
             "Backup-Test: vierteljährlich, protokolliert in FORM-BK-2025."),
            ("Sicherheitsaudit",
             "Halbjährliches Netzwerksicherheitsaudit durch MA-025 Oliver Großmann. "
             "Penetrationstest: jährlich durch externen Dienstleister. "
             "Ergebnisse: vertraulich, aufzubewahren 5 Jahre."),
        ]),
        ("Anhang", [
            ("Technische Daten im Überblick",
             "Gehäuse: 19\" / 2U, Schutzart IP67 (bei montierten Steckverbindern). "
             "Betriebstemperatur: −20 °C bis +60 °C. Lagertemperatur: −40 °C bis +70 °C. "
             "Relative Luftfeuchtigkeit: 5–95 % (nicht kondensierend). Gewicht: 2,8 kg."),
            ("Zulassungen und Zertifikate",
             "CE-Kennzeichnung, UL-Listing (Nordamerika), RoHS-konform, REACH-konform. "
             "ISO 9001:2015 (Musterbau GmbH, Dortmund, HRB 28471 Dortmund). "
             "ATEX/IECEx auf Anfrage."),
            ("Abkürzungsverzeichnis",
             "CVE: Zentrale Verarbeitungseinheit. EMV: Elektromagnetische Verträglichkeit. "
             "FTE: Vollzeitäquivalent. KOM: Kommunikationsmodul. MQTT: Message Queuing Telemetry Transport. "
             "OPC-UA: Open Platform Communications Unified Architecture. PIMS: Prozessintegriertes Monitoring-System. "
             "QES: Qualifizierte elektronische Signatur. RPO: Recovery Point Objective. RTO: Recovery Time Objective. "
             "SIM: Sensorinterface-Modul. SNMP: Simple Network Management Protocol. USV: Unterbrechungsfreie Stromversorgung."),
            ("Ansprechpartner",
             f"{FIRMA} | {SITZ} | HRB: {HRB} | Steuer-ID: {STEUER_ID}\n"
             f"Technisch: Tobias Wernecke (MA-011), Engineering-Leiter\n"
             f"Qualität: Sandra Kemper (MA-019), Qualitätssicherung\n"
             f"IT/Administration: Oliver Großmann (MA-025)\n"
             f"Geschäftsführung: {GF} (MA-001)"),
            ("Revisionshistorie dieses Handbuchs",
             "V1.0 (2023-01): Erstausgabe für PIMS-A 1.x. "
             "V2.0 (2024-06): 24-Kanal-Erweiterung, Abschnitt Kalibrierung neu. "
             "V3.0 (2025-01): Kavitationsschutzprotokoll, Drewermann-Verfahren, Fensterprotokoll 24-Kanal. "
             "V3.1 (2025-09): Nach ISO-9001-Rezertifizierung (E-04) aktualisiert. "
             "V3.2 (2025-12): Redaktionelle Überarbeitung."),
        ]),
    ]

    story = []

    # Cover
    story.append(Spacer(1, 2 * cm))
    story.append(Paragraph(
        f"<font color='#213452'><b>Technisches Betriebshandbuch</b></font>",
        sty("covH", fontSize=24, leading=30, fontName="Helvetica-Bold")))
    story.append(Paragraph(
        "PIMS-A Basismodul",
        sty("covH2", fontSize=18, leading=22, textColor=ORANGE_RL, fontName="Helvetica-Bold")))
    story.append(Paragraph(
        f"{FIRMA} | {SITZ} | {HRB}",
        sty("covS", fontSize=10, textColor=HexColor("#555555"), spaceBefore=6)))
    story.append(Paragraph(
        "Dokumentversion 3.2 | Januar 2026 | Vertraulich — nur für internen Gebrauch",
        sty("covS2", fontSize=9, textColor=HexColor("#888888"), spaceAfter=20)))
    story.append(PageBreak())

    for ch_idx, (chapter_title, sections) in enumerate(chapters, start=1):
        story.append(Paragraph(f"Kapitel {ch_idx}: {chapter_title}", H1))
        for sec_title, sec_body in sections:
            story.append(Paragraph(sec_title, H2))
            # Emit 4–5 paragraphs per section to generate page volume
            story.append(Paragraph(sec_body, BODY))
            # Expand with supplementary paragraphs to increase page count to ≥50
            story.append(Paragraph(
                f"Detaillierte Erläuterung zu '{sec_title}': Die ordnungsgemäße Durchführung "
                f"aller hier beschriebenen Maßnahmen liegt in der Verantwortung des jeweiligen "
                f"Betreibers. Bei Fragen wenden Sie sich an die zuständige Fachabteilung "
                f"oder direkt an Tobias Wernecke (MA-011), Engineering-Leiter, der {FIRMA}, "
                f"{SITZ} ({HRB}).",
                NOTE))
            story.append(Paragraph(
                f"Verfahrensanweisung: Vor Beginn aller Arbeiten an diesem Abschnitt ist "
                f"sicherzustellen, dass sämtliche Sicherheitsvorschriften eingehalten werden. "
                f"Die Verantwortung für die korrekte Umsetzung obliegt dem nominierten "
                f"Inbetriebnahmetechniker. Jede Abweichung vom beschriebenen Verfahren ist "
                f"schriftlich zu dokumentieren und dem Projektleiter zu melden.",
                NOTE))
            story.append(Paragraph(
                f"Qualitätssicherungshinweis: Alle durchgeführten Arbeiten sind gemäß "
                f"ISO 9001:2015 (Zertifikat der {FIRMA}, erteilt 22.05.2025, E-04) zu "
                f"dokumentieren. Zuständig für Qualitätssicherung: MA-019 Sandra Kemper. "
                f"Protokolle sind 10 Jahre aufzubewahren.",
                NOTE))
            story.append(Paragraph(
                f"Hinweis: Alle Angaben basieren auf dem Stand der Technik zum Zeitpunkt "
                f"der Drucklegung dieses Handbuchs (Version 3.2, Januar 2026). "
                f"Technische Änderungen vorbehalten. Maßgeblich ist stets die jeweils "
                f"gültige Version auf dem Kundenportal der {FIRMA}.",
                NOTE))
            story.append(Paragraph(
                f"Sicherheitsrelevante Informationen zu diesem Abschnitt: "
                f"PIMS-A ist gemäß IEC 61508 SIL 2 klassifiziert. "
                f"Abweichungen von dieser Anleitung können die Produktsicherheit beeinträchtigen "
                f"und zum Erlöschen der Gewährleistung führen. "
                f"Kontakt Geschäftsführung: Werner Kahlert (MA-001), {FIRMA}.",
                NOTE))
            story.append(Paragraph(
                f"Ergänzende Betriebsinformation: Der Betrieb von PIMS-A in Umgebungen "
                f"mit erhöhter elektromagnetischer Störstrahlung erfordert zusätzliche "
                f"Abschirmmaßnahmen gemäß EN 61000-6-2. Bei Rückfragen zur Systemintegration "
                f"steht MA-012 Sven Unterberg (Senior-Ingenieur, ENG) als Ansprechpartner "
                f"zur Verfügung.",
                NOTE))
            story.append(Paragraph(
                f"Referenzdokumente: Technische Spezifikation PIMS-A (Dok.-Nr. MB-PIMS-A-TS-2025), "
                f"Wartungsplan FORM-WP-2025-A, Kalibrierformular FORM-KAL-2025, "
                f"Kavitationsschutz-Formular FORM-KSP-2025, Backup-Protokoll FORM-BK-2025. "
                f"Alle Formulare sind im internen Dokumentenmanagementsystem hinterlegt.",
                NOTE))
            story.append(Spacer(1, 0.15 * cm))
        story.append(PageBreak())

    class HeaderFooterCanvas(rl_canvas.Canvas):
        def __init__(self, filename, **kwargs):
            super().__init__(filename, **kwargs)
            self._page_num = 0

        def showPage(self):
            self._page_num += 1
            self._draw_hf()
            super().showPage()

        def save(self):
            self._page_num += 1
            self._draw_hf()
            super().save()

        def _draw_hf(self):
            _draw_page_header(self, PAGE_W, PAGE_H, doc_title)
            _draw_page_footer(self, PAGE_W, self._page_num)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=A4,
        leftMargin=MARGIN, rightMargin=MARGIN,
        topMargin=3.5 * cm, bottomMargin=2.5 * cm,
        title=doc_title,
        author=FIRMA,
    )
    doc.build(story, canvasmaker=HeaderFooterCanvas)


# ===========================================================================
# 5. dedup/sample.pdf
# ===========================================================================

def build_dedup_sample(out_path: Path) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.colors import HexColor
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.pdfgen import canvas as rl_canvas

    PAGE_W, PAGE_H = A4
    MARGIN = 3.0 * cm
    NAVY_RL = HexColor("#213452")

    styles = getSampleStyleSheet()
    BODY = ParagraphStyle("dedup_body", parent=styles["Normal"], fontSize=10, leading=14, spaceAfter=8)
    H1 = ParagraphStyle("dedup_h1", parent=styles["Normal"], fontSize=14, leading=18,
                         fontName="Helvetica-Bold", textColor=NAVY_RL, spaceAfter=10, spaceBefore=12)

    story = []
    story.append(Paragraph(f"{FIRMA} — Referenz-Dokument (Dedup-Test)", H1))
    story.append(Paragraph(
        f"Dieses Dokument enthält identifizierbare Musterbau-Kerndaten und dient als "
        f"Referenz für den Inhalts-Hash-Dedup-Test des RAG-Systems. "
        f"Es wird zweimal hochgeladen; das System soll beim zweiten Upload keinen neuen "
        f"Chunk erzeugen.",
        BODY))
    story.append(Paragraph(
        f"Firmenname: {FIRMA} | Handelsregisternummer: {HRB} | Sitz: {SITZ} | "
        f"Gründungsjahr: {GRUENDUNGSJAHR} | Steuer-ID: {STEUER_ID} | "
        f"Geschäftsführer: {GF} (MA-001) | CFO: {CFO} (MA-002, seit {CFO_START}).",
        BODY))
    story.append(Paragraph(
        f"Umsatz 2025: {UMSATZ_2025:,} € | Jahresüberschuss: {JAHRESUEBERSCHUSS:,} € | "
        f"Bilanzsumme: {BILANZ_AKTIVA:,} € | Mitarbeitende (FTE): {FTE}.".replace(",", "."),
        BODY))
    story.append(Paragraph(
        f"Produktlinien: PIMS-A (48.900 €), PIMS-B (18.500 €), PIMS-C (12.200 €), "
        f"PIMS-D (29.700 €), PIMS-E (9.800 €). Größter Kunde: KD-007 Bramkamp "
        f"Industrietechnik GmbH (742.000 € p.a.), betreut von Dirk Hammerschmidt (MA-007).",
        BODY))
    story.append(Paragraph(
        f"Schlüsselereignisse 2025: E-01 Strategieklausur (14.01.), E-02 CFO-Wechsel "
        f"(01.04.), E-03 Großauftrag KD-007 (18.04.), E-04 ISO-9001-Audit (22.05.), "
        f"E-05 {TERM_SAUERLANDSTERN} (10.07.), E-06 Q3-Review (15.09.), "
        f"E-07 Produktionsoptimierung (03.11.), E-08 Jahresabschluss (09.12.).",
        BODY))

    class HFC(rl_canvas.Canvas):
        def __init__(self, filename, **kwargs):
            super().__init__(filename, **kwargs)
            self._pn = 0

        def showPage(self):
            self._pn += 1
            _draw_page_header(self, PAGE_W, PAGE_H, f"{FIRMA} — Dedup-Referenz")
            _draw_page_footer(self, PAGE_W, self._pn)
            super().showPage()

        def save(self):
            self._pn += 1
            _draw_page_header(self, PAGE_W, PAGE_H, f"{FIRMA} — Dedup-Referenz")
            _draw_page_footer(self, PAGE_W, self._pn)
            super().save()

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(str(out_path), pagesize=A4,
                            leftMargin=MARGIN, rightMargin=MARGIN,
                            topMargin=3.5 * cm, bottomMargin=2.5 * cm)
    doc.build(story, canvasmaker=HFC)


# ===========================================================================
# 6. phash/logo_repeating.pdf
# ===========================================================================

def build_phash_logo_repeating(out_path: Path) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.colors import HexColor, white
    from reportlab.lib.units import cm
    from reportlab.graphics.shapes import Drawing, Rect, String
    from reportlab.graphics.charts.barcharts import VerticalBarChart
    from reportlab.graphics.charts.piecharts import Pie
    from reportlab.pdfgen import canvas as rl_canvas

    PAGE_W, PAGE_H = A4
    NAVY_RL = HexColor("#213452")
    ORANGE_RL = HexColor("#ee7f00")

    # Logo is always drawn at the same centre position on each page.
    LOGO_X = (PAGE_W - 200) / 2  # centred
    LOGO_Y = PAGE_H / 2 + 30
    LOGO_W = 200
    LOGO_H = 50

    def draw_centred_logo(c):
        """Identical logo on all 3 pages — same position, size, colour."""
        c.saveState()
        # Navy rect
        c.setFillColor(NAVY_RL)
        c.rect(LOGO_X, LOGO_Y, LOGO_W, LOGO_H, fill=1, stroke=0)
        # Orange left stripe
        c.setFillColor(ORANGE_RL)
        c.rect(LOGO_X, LOGO_Y, 14, LOGO_H, fill=1, stroke=0)
        # White text
        c.setFillColor(white)
        c.setFont("Helvetica-Bold", 20)
        c.drawString(LOGO_X + 22, LOGO_Y + 17, "M GmbH")
        # Subtitle
        c.setFont("Helvetica", 9)
        c.drawString(LOGO_X + 22, LOGO_Y + 5, "Musterbau GmbH — Dortmund")
        c.restoreState()

    # Unique figures: bar chart (page 2), pie chart (page 3)
    def draw_bar_chart(c):
        from reportlab.graphics import renderPDF
        d = Drawing(260, 140)
        bc = VerticalBarChart()
        bc.x = 30
        bc.y = 20
        bc.width = 200
        bc.height = 100
        bc.data = [[742, 683, 561, 437]]
        bc.categoryAxis.categoryNames = ["KD-007", "KD-034", "KD-021", "KD-015"]
        bc.bars[0].fillColor = ORANGE_RL
        bc.valueAxis.valueMin = 0
        bc.valueAxis.valueMax = 800
        bc.valueAxis.valueStep = 200
        bc.valueAxis.labelTextFormat = "%d T€"
        d.add(bc)
        lbl = String(130, 130, "Großkunden 2025 (T€)", textAnchor="middle",
                     fontSize=9, fillColor=NAVY_RL, fontName="Helvetica-Bold")
        d.add(lbl)
        chart_x = (PAGE_W - 260) / 2
        chart_y = LOGO_Y - 200
        renderPDF.draw(d, c, chart_x, chart_y)

    def draw_pie_chart(c):
        from reportlab.graphics import renderPDF
        d = Drawing(260, 180)
        pie = Pie()
        pie.x = 80
        pie.y = 20
        pie.width = 120
        pie.height = 120
        pie.data = [UMSATZERLOESE, SONST_ERTRAGE]
        pie.labels = ["Umsatzerlöse", "Sonst. Erträge"]
        pie.slices[0].fillColor = NAVY_RL
        pie.slices[1].fillColor = ORANGE_RL
        d.add(pie)
        lbl = String(130, 165, "Gesamtleistung 2025", textAnchor="middle",
                     fontSize=9, fillColor=NAVY_RL, fontName="Helvetica-Bold")
        d.add(lbl)
        chart_x = (PAGE_W - 260) / 2
        chart_y = LOGO_Y - 230
        renderPDF.draw(d, c, chart_x, chart_y)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    c = rl_canvas.Canvas(str(out_path), pagesize=A4)

    for page_num in range(1, 4):
        # Standard page header (logo in top-right via header helper)
        _draw_page_header(c, PAGE_W, PAGE_H, f"{FIRMA} — Logo-Wiederholungstest (pHash)")
        _draw_page_footer(c, PAGE_W, page_num, 3)

        # Page label
        c.setFont("Helvetica-Bold", 12)
        c.setFillColor(NAVY_RL)
        c.drawCentredString(PAGE_W / 2, PAGE_H - 90, f"Seite {page_num} — Logo-Wiederholungstest")

        # SAME centred logo on every page (same geometry → same pHash region)
        draw_centred_logo(c)

        # Page label inside logo area
        c.setFont("Helvetica", 9)
        c.setFillColor(HexColor("#555555"))
        c.drawCentredString(PAGE_W / 2, LOGO_Y - 20,
                            "Dieses Logo erscheint auf jeder Seite identisch.")

        # Unique figure on pages 2 and 3
        if page_num == 2:
            draw_bar_chart(c)
            c.setFont("Helvetica", 9)
            c.setFillColor(HexColor("#555555"))
            c.drawCentredString(PAGE_W / 2, LOGO_Y - 215,
                                "Abb. 1 (einmalig): Großkundenumsätze 2025")
        elif page_num == 3:
            draw_pie_chart(c)
            c.setFont("Helvetica", 9)
            c.setFillColor(HexColor("#555555"))
            c.drawCentredString(PAGE_W / 2, LOGO_Y - 250,
                                "Abb. 2 (einmalig): Gesamtleistung-Aufteilung 2025")

        c.showPage()

    c.save()


# ===========================================================================
# 7. multidoc_a/appendix.pdf
# ===========================================================================

def build_multidoc_appendix(out_path: Path) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.colors import HexColor, white
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, HRFlowable,
    )
    from reportlab.pdfgen import canvas as rl_canvas

    PAGE_W, PAGE_H = A4
    MARGIN = 2.5 * cm
    NAVY_RL = HexColor("#213452")
    ORANGE_RL = HexColor("#ee7f00")

    styles = getSampleStyleSheet()

    def sty(name, **kw):
        return ParagraphStyle(name, parent=styles["Normal"], **kw)

    H1 = sty("H1ap", fontSize=14, textColor=NAVY_RL, spaceBefore=16, spaceAfter=6,
             fontName="Helvetica-Bold", leading=18)
    H2 = sty("H2ap", fontSize=11, textColor=NAVY_RL, spaceBefore=10, spaceAfter=4,
             fontName="Helvetica-Bold", leading=14)
    BODY = sty("Bodyap", fontSize=10, leading=14, spaceAfter=7)
    SMALL = sty("Smallap", fontSize=9, leading=12, spaceAfter=5,
                textColor=HexColor("#444444"))

    doc_title = f"{FIRMA} — Anhang Großauftrag KD-007 (Bramkamp Industrietechnik GmbH)"

    story = []
    story.append(Spacer(1, 1 * cm))

    # Cover block
    story.append(Paragraph("Anhang zum Großauftrag KD-007", H1))
    story.append(Paragraph(
        "Bramkamp Industrietechnik GmbH | PIMS-A + PIMS-D Rahmenvertrag",
        sty("apSub", fontSize=12, textColor=ORANGE_RL, fontName="Helvetica-Bold", spaceAfter=6)))
    story.append(Paragraph(
        f"Auftraggeber: KD-007 Bramkamp Industrietechnik GmbH, Hamm (NRW) | "
        f"Auftragnehmer: {FIRMA}, {SITZ} | Key-Account: Dirk Hammerschmidt (MA-007) | "
        f"Vertragsvolumen: 742.000 € p.a., Laufzeit 3 Jahre (ab 18.04.2025, E-03)",
        BODY))
    story.append(HRFlowable(width="100%", thickness=1.5, color=ORANGE_RL))
    story.append(Spacer(1, 0.3 * cm))

    # Section 1: Kostenzusammenstellung
    story.append(Paragraph("1. Kostenzusammenstellung", H1))
    story.append(Paragraph(
        "Die folgende Tabelle zeigt die vertraglich vereinbarten Leistungspositionen "
        "und Preise für den Rahmenvertrag mit Bramkamp Industrietechnik GmbH (KD-007).",
        BODY))

    cost_data = [
        ["Pos.", "Leistung", "Menge", "Einzelpreis (€)", "Gesamtpreis (€)"],
        ["1", "PIMS-A Basismodul (Kernplattform)", "8", "48.900", "391.200"],
        ["2", "PIMS-D Predictive-Maintenance-Modul", "8", "29.700", "237.600"],
        ["3", "Inbetriebnahme und Konfiguration (Tagessatz)", "12", "1.850", "22.200"],
        ["4", "Schulung Kundenpersonal (2 Tage)", "2", "3.500", "7.000"],
        ["5", "Jahreswartungsvertrag (3 Jahre × 28.000 €)", "3", "28.000", "84.000"],
        ["", "Gesamtvolumen (3 Jahre)", "", "", "742.000"],
    ]

    cost_table = Table(cost_data, colWidths=[1 * cm, 7 * cm, 1.5 * cm, 3 * cm, 3 * cm])
    cost_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), NAVY_RL),
        ("TEXTCOLOR", (0, 0), (-1, 0), white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTNAME", (0, -1), (-1, -1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("ROWBACKGROUNDS", (0, 1), (-1, -2), [HexColor("#f5f5f5"), white]),
        ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#cccccc")),
        ("ALIGN", (2, 0), (-1, -1), "RIGHT"),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BACKGROUND", (0, -1), (-1, -1), HexColor("#EE7F00")),
        ("TEXTCOLOR", (0, -1), (-1, -1), white),
    ]))
    story.append(cost_table)
    story.append(Spacer(1, 0.3 * cm))
    story.append(Paragraph(
        "Alle Preise verstehen sich netto zzgl. gesetzlicher MwSt. "
        "Listenpreise PIMS-A: 48.900 €, PIMS-D: 29.700 € (Abschnitt 5, MUSTERBAU-Spezifikation). "
        "Vertragsrabatt: gemäß Rahmenvertragskonditionen KD-007.",
        SMALL))

    story.append(PageBreak())

    # Section 2: Abnahmekriterien
    story.append(Paragraph("2. Abnahmekriterien", H1))
    story.append(Paragraph(
        "Die folgenden Kriterien wurden zwischen Bramkamp Industrietechnik GmbH (KD-007) "
        "und der Musterbau GmbH schriftlich vereinbart und gelten als verbindliche "
        "Abnahmevoraussetzung:",
        BODY))

    crit_data = [
        ["Nr.", "Kriterium", "Prüfmethode", "Verantwortlich", "Status"],
        ["AC-01", "PIMS-A: alle 24 Kanäle kalibriert (Drewermann-Verfahren)",
         "Kalibrierzertifikat", "MA-019 S. Kemper", "Offen"],
        ["AC-02", "PIMS-D: Prognosegüte ≥ 85 % (Testdatensatz 6 Monate)",
         "Backtesting-Report", "MA-011 T. Wernecke", "Offen"],
        ["AC-03", "Systemverfügbarkeit ≥ 99,5 % über 30-Tage-Testbetrieb",
         "Monitoring-Log", "MA-025 O. Großmann", "Offen"],
        ["AC-04", "Schulungsnachweis für 6 Kundentechniker",
         "Teilnehmerliste", "MA-007 D. Hammerschmidt", "Offen"],
        ["AC-05", "Übergabe vollständige Dokumentation (DE/EN)",
         "Dokumentenprüfung", "MA-011 T. Wernecke", "Offen"],
    ]

    crit_table = Table(crit_data, colWidths=[1.2 * cm, 5.5 * cm, 3 * cm, 3.5 * cm, 2 * cm])
    crit_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), NAVY_RL),
        ("TEXTCOLOR", (0, 0), (-1, 0), white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [HexColor("#f5f5f5"), white]),
        ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#cccccc")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(crit_table)
    story.append(Spacer(1, 0.3 * cm))
    story.append(Paragraph(
        "Abnahme erfolgt durch schriftliche Bestätigung beider Vertragsparteien. "
        "Abnahmedatum: spätestens 30 Tage nach Abschluss des Testbetriebs. "
        "Ansprechpartner Musterbau: Dirk Hammerschmidt (MA-007), Werner Kahlert (MA-001).",
        SMALL))

    story.append(PageBreak())

    # Section 3: Meilensteinplan
    story.append(Paragraph("3. Meilensteinplan", H1))
    story.append(Paragraph(
        "Der folgende Meilensteinplan gilt für die Erstlieferung und Inbetriebnahme "
        "der 8 PIMS-A- und 8 PIMS-D-Systeme bei Bramkamp Industrietechnik GmbH (KD-007).",
        BODY))

    ms_data = [
        ["MS", "Meilenstein", "Termin", "Verantwortlich"],
        ["MS-1", "Vertragsunterzeichnung und Anzahlung (30 %)", "18.04.2025 (E-03)", "MA-007"],
        ["MS-2", "Lieferung Hardware und Lizenzaktivierung", "30.06.2025", "MA-011, MA-020"],
        ["MS-3", "Inbetriebnahme und Kalibrierung (Drewermann-Verfahren)", "31.08.2025", "MA-011"],
        ["MS-4", "Testbetrieb 30 Tage (Abnahmekriterien AC-01 bis AC-05)", "30.09.2025", "MA-019"],
        ["MS-5", "Schulung Kundenpersonal", "15.10.2025", "MA-007"],
        ["MS-6", "Abnahme und Restrechnung (70 %)", "31.10.2025", "MA-007, MA-002"],
        ["MS-7", "Jahreswartungsvertrag aktiv", "01.11.2025", "MA-011"],
    ]

    ms_table = Table(ms_data, colWidths=[1.2 * cm, 7 * cm, 3.5 * cm, 3.5 * cm])
    ms_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), NAVY_RL),
        ("TEXTCOLOR", (0, 0), (-1, 0), white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [HexColor("#f5f5f5"), white]),
        ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#cccccc")),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(ms_table)
    story.append(Spacer(1, 0.3 * cm))
    story.append(Paragraph(
        f"Gesamtprojektleitung: Dirk Hammerschmidt (MA-007). "
        f"Technische Leitung: Tobias Wernecke (MA-011). "
        f"Finanzielle Freigaben: {CFO} (MA-002). "
        f"Dieses Dokument ist Teil der Vertragsdokumentation zu Ereignis E-03 "
        f"(18.04.2025, Hamm — Bramkamp-Zentrale).",
        SMALL))

    story.append(PageBreak())

    # Section 4: Unterzeichnung
    story.append(Paragraph("4. Freigabe und Unterzeichnung", H1))
    story.append(Paragraph(
        "Dieses Anhangdokument wurde von beiden Vertragsparteien geprüft und freigegeben.",
        BODY))

    sign_data = [
        ["Partei", "Name", "Funktion", "Datum", "Unterschrift"],
        ["Musterbau GmbH", "Werner Kahlert (MA-001)", "Geschäftsführer", "18.04.2025", "___________"],
        ["Musterbau GmbH", "Dirk Hammerschmidt (MA-007)", "Key-Account", "18.04.2025", "___________"],
        ["Bramkamp Industrietechnik GmbH", "N.N.", "Geschäftsführer", "18.04.2025", "___________"],
    ]

    sign_table = Table(sign_data, colWidths=[3.5 * cm, 3.5 * cm, 3 * cm, 2.5 * cm, 2.5 * cm])
    sign_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), NAVY_RL),
        ("TEXTCOLOR", (0, 0), (-1, 0), white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [HexColor("#f5f5f5"), white]),
        ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#cccccc")),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(sign_table)
    story.append(Spacer(1, 0.5 * cm))
    story.append(Paragraph(
        f"Musterbau GmbH | {SITZ} | {HRB} | Steuer-ID: {STEUER_ID}",
        sty("foot_ap", fontSize=8, textColor=HexColor("#888888"))))

    class HFC(rl_canvas.Canvas):
        def __init__(self, filename, **kwargs):
            super().__init__(filename, **kwargs)
            self._pn = 0

        def showPage(self):
            self._pn += 1
            _draw_page_header(self, PAGE_W, PAGE_H, doc_title)
            _draw_page_footer(self, PAGE_W, self._pn)
            super().showPage()

        def save(self):
            self._pn += 1
            _draw_page_header(self, PAGE_W, PAGE_H, doc_title)
            _draw_page_footer(self, PAGE_W, self._pn)
            super().save()

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(str(out_path), pagesize=A4,
                            leftMargin=MARGIN, rightMargin=MARGIN,
                            topMargin=3.5 * cm, bottomMargin=2.5 * cm,
                            title=doc_title, author=FIRMA)
    doc.build(story, canvasmaker=HFC)


# ===========================================================================
# Legacy .xls (BIFF8) builder — minimal Bilanz sheet via xlwt
# ===========================================================================

def build_finanzen_legacy_xls(out_path: Path) -> None:
    """One-sheet Bilanz_Alt export in the legacy .xls (BIFF8) format.

    Mirrors a subset of finanzen_2025.xlsx so the same Aktiva totals roll
    through, but uses xlwt (limited styling) and lives in a single sheet —
    a realistic shape for a legacy-format export from a pre-2007 toolchain.
    Exercises the .xls extraction path; not styled for human review.
    """
    import xlwt

    wb = xlwt.Workbook(encoding="utf-8")
    ws = wb.add_sheet("Bilanz_Alt")

    header_style = xlwt.easyxf(
        "font: bold on, colour white; "
        "pattern: pattern solid, fore_colour dark_blue; "
        "align: horiz center"
    )
    total_style = xlwt.easyxf(
        "font: bold on, colour white; "
        "pattern: pattern solid, fore_colour orange",
        num_format_str='#,##0 "€"',
    )
    eur_style = xlwt.easyxf(num_format_str='#,##0 "€"')
    italic_style = xlwt.easyxf("font: italic on")

    ws.col(0).width = 16000
    ws.col(1).width = 6000

    ws.write(0, 0, f"{FIRMA} — Bilanz zum 31.12.2025 (Altformat-Export)")
    ws.write(1, 0, f"HRB: {HRB} | Steuer-ID: {STEUER_ID}", italic_style)
    ws.write(2, 0, f"Verantwortlich: {CFO} (MA-002); Stichtag-Bewertung im Zuge des CFO-Wechsels E-02 am {CFO_START}.", italic_style)

    rows = [
        ("Position", "Betrag (€)", header_style, header_style),
        ("AKTIVA", "", header_style, header_style),
        ("Anlagevermögen gesamt", ANLAGEVERMOEGEN, None, eur_style),
        ("  — Sachanlagen", SACHANLAGEN, None, eur_style),
        ("  — Immaterielle Vermögenswerte", IMMAT_VERMOEGEN, None, eur_style),
        ("Umlaufvermögen gesamt", UMLAUFVERMOEGEN, None, eur_style),
        ("  — Vorräte", VORRAETE, None, eur_style),
        ("  — Forderungen aus L+L", FORDERUNGEN, None, eur_style),
        ("  — Kassenbestand und Bankguthaben", KASSENBESTAND, None, eur_style),
        ("Bilanzsumme Aktiva", BILANZ_AKTIVA, total_style, total_style),
    ]
    r = 4
    for left, right, ls, rs in rows:
        if ls is None:
            ws.write(r, 0, left)
        else:
            ws.write(r, 0, left, ls)
        if right == "":
            pass
        elif rs is None:
            ws.write(r, 1, right)
        else:
            ws.write(r, 1, right, rs)
        r += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(out_path))


# ===========================================================================
# Rheintal corpus builders — Kunstakademie Rheintal e. V. / Rheintal Akademie gGmbH
# Source of truth: docs/tests/phase3/corpus/rheintal/RHEINTAL.md (frozen v1.0)
# All values are hardcoded inline from that document; do NOT mix with Musterbau constants.
# ===========================================================================


def _draw_kar_logo(canvas, x: float, y: float, width: float = 110, height: float = 28) -> None:
    """Stylized KAR logo box: navy rectangle + orange left stripe + white 'KAR' text."""
    from reportlab.lib.colors import HexColor
    navy = HexColor("#213452")
    orange = HexColor("#ee7f00")
    white = HexColor("#ffffff")

    canvas.saveState()
    canvas.setFillColor(navy)
    canvas.rect(x, y, width, height, fill=1, stroke=0)
    canvas.setFillColor(orange)
    canvas.rect(x, y, 7, height, fill=1, stroke=0)
    canvas.setFillColor(white)
    canvas.setFont("Helvetica-Bold", 11)
    canvas.drawString(x + 13, y + 9, "KAR")
    canvas.restoreState()


def _draw_kar_page_header(canvas, page_width: float, page_height: float, title_text: str) -> None:
    """Page header with KAR logo top-right and document title top-left."""
    from reportlab.lib.colors import HexColor
    navy = HexColor("#213452")

    _draw_kar_logo(canvas, page_width - 145, page_height - 48)
    canvas.saveState()
    canvas.setStrokeColor(navy)
    canvas.setLineWidth(1.5)
    canvas.line(40, page_height - 58, page_width - 40, page_height - 58)
    canvas.setFillColor(navy)
    canvas.setFont("Helvetica-Bold", 9)
    canvas.drawString(40, page_height - 48, title_text)
    canvas.restoreState()


def _draw_kar_page_footer(canvas, page_width: float, page_num: int) -> None:
    """Page footer: thin navy rule, org name left, page number right."""
    from reportlab.lib.colors import HexColor
    grey = HexColor("#555555")
    canvas.saveState()
    canvas.setStrokeColor(HexColor("#213452"))
    canvas.setLineWidth(0.5)
    canvas.line(40, 38, page_width - 40, 38)
    canvas.setFillColor(grey)
    canvas.setFont("Helvetica", 8)
    canvas.drawString(40, 26, "Kunstakademie Rheintal e. V. / Rheintal Akademie gGmbH | VR 4312 Freiburg | 2025")
    canvas.drawRightString(page_width - 40, 26, f"Seite {page_num}")
    canvas.restoreState()


# ---------------------------------------------------------------------------
# 1. taetigkeitsbericht_2025.pdf
# ---------------------------------------------------------------------------

def build_taetigkeitsbericht(out_path: Path) -> None:
    """4-page Tätigkeitsbericht 2025 der Kunstakademie Rheintal e. V. + Rheintal Akademie gGmbH.

    BM25 target terms: Werkförderverfahren (≥ 2×), Druckgrafik-Residenz (≥ 2×).
    References: PER-001, PER-002, PER-003, PER-009, PER-011, KL-01.
    Events: EV-01, EV-03, EV-05, EV-08.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.colors import HexColor, white
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak,
        HRFlowable,
    )
    from reportlab.pdfgen import canvas as rl_canvas

    PAGE_W, PAGE_H = A4
    MARGIN = 2.5 * cm
    NAVY_RL = HexColor("#213452")
    ORANGE_RL = HexColor("#ee7f00")

    styles = getSampleStyleSheet()

    def sty(name, **kw):
        base = kw.pop("base", "Normal")
        return ParagraphStyle(name, parent=styles[base], **kw)

    TITLE = sty("KARTitle", fontSize=22, textColor=NAVY_RL, spaceBefore=0, spaceAfter=8,
                fontName="Helvetica-Bold", leading=28)
    SUBTITLE = sty("KARSub", fontSize=14, textColor=ORANGE_RL, spaceBefore=4, spaceAfter=12,
                   fontName="Helvetica-Bold")
    H1 = sty("KARh1", fontSize=14, textColor=NAVY_RL, spaceBefore=16, spaceAfter=6,
             fontName="Helvetica-Bold")
    H2 = sty("KARh2", fontSize=11, textColor=NAVY_RL, spaceBefore=10, spaceAfter=4,
             fontName="Helvetica-Bold")
    BODY = sty("KARBody", fontSize=10, leading=14, spaceBefore=4, spaceAfter=4)
    SMALL = sty("KARSmall", fontSize=8, leading=11, textColor=HexColor("#555555"))

    doc_title = "Tätigkeitsbericht 2025 — Kunstakademie Rheintal e. V."

    story = []

    # ── Page 1: Title + Organisation Profile ────────────────────────────────
    story.append(Spacer(1, 0.4 * cm))
    story.append(Paragraph("Tätigkeitsbericht 2025", TITLE))
    story.append(Paragraph("Kunstakademie Rheintal e. V. und Rheintal Akademie gemeinnützige GmbH", SUBTITLE))
    story.append(HRFlowable(width="100%", thickness=2, color=ORANGE_RL, spaceAfter=12))

    story.append(Paragraph("1. Organisationsprofil", H1))
    story.append(Paragraph(
        "Die Kunstakademie Rheintal e. V. wurde 1997 gegründet und hat ihren Sitz in "
        "Freiburg im Breisgau, Baden-Württemberg (Vereinsregisternummer VR 4312 Freiburg im Breisgau). "
        "Der Verein ist als gemeinnützige Organisation nach §52 AO von der Körperschaftsteuer befreit "
        "und zählt zum Stichtag 31. Dezember 2025 insgesamt 318 ordentliche Mitglieder.",
        BODY))
    story.append(Paragraph(
        "Die operative Kursabwicklung sowie alle Vertragsabschlüsse mit Honorarkräften werden über die "
        "hundertprozentige Tochtergesellschaft Rheintal Akademie gemeinnützige GmbH (Stammkapital 25.000 €) "
        "abgewickelt. Geschäftsführerin der gGmbH ist Dr. Margit Feuerbach (PER-002, im Amt seit 01.09.2008). "
        "Vereinsvorsitzende ist Renate Quast (PER-001). Programmdirektor ist Lukas Endres (PER-003).",
        BODY))
    story.append(Paragraph(
        "Im Berichtsjahr 2025 wurden 87 Kurse in elf Fachbereichen angeboten, betreut von 34 aktiven "
        "Honorar-Kursleitern (KL-01 bis KL-34) und 21 Festangestellten (PER-001 bis PER-021). "
        "Die geschätzte Gesamtzahl der Kursteilnehmer lag bei rund 1.200 Personen. "
        "Ausstellungskoordinatorin Franziska Oppelt (PER-009) betreute sechs öffentliche Ausstellungen. "
        "Urte Hamann (PER-011) verantwortete die Beantragung und Abwicklung aller Drittmittel.",
        BODY))

    org_data = [
        ["Feld", "Wert"],
        ["Organisation", "Kunstakademie Rheintal e. V."],
        ["Vereinsregisternummer", "VR 4312 Freiburg im Breisgau"],
        ["Gründungsjahr", "1997"],
        ["Sitz", "Freiburg im Breisgau, Baden-Württemberg"],
        ["Rechtsform Träger", "Eingetragener Verein (e. V.)"],
        ["100%-Tochter", "Rheintal Akademie gemeinnützige GmbH"],
        ["gGmbH Stammkapital", "25.000 €"],
        ["Steuerlicher Status", "Gemeinnützig nach §52 AO, steuerbefreit"],
        ["Jahresbudget 2025", "1.240.000 € (Einnahmen = Ausgaben; kein Jahresüberschuss)"],
        ["Vereinsvorsitzende", "Renate Quast (PER-001)"],
        ["Geschäftsführerin gGmbH", "Dr. Margit Feuerbach (PER-002, seit 01.09.2008)"],
    ]

    org_table = Table(org_data, colWidths=[6 * cm, PAGE_W - 2 * MARGIN - 6 * cm])
    org_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), NAVY_RL),
        ("TEXTCOLOR", (0, 0), (-1, 0), white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [HexColor("#f5f5f5"), white]),
        ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#cccccc")),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(Spacer(1, 0.4 * cm))
    story.append(org_table)

    story.append(PageBreak())

    # ── Page 2: Programmhöhepunkte 2025 ─────────────────────────────────────
    story.append(Paragraph("2. Programmhöhepunkte 2025", H1))

    story.append(Paragraph("2.1 Druckgrafik-Residenz (EV-01)", H2))
    story.append(Paragraph(
        "Auf der Jahresauftaktklausur des Vorstands am 15. Januar 2025 (EV-01) beschlossen "
        "Renate Quast (PER-001), Dr. Margit Feuerbach (PER-002), Lukas Endres (PER-003) und "
        "Urte Hamann (PER-011) die Einrichtung einer Druckgrafik-Residenz an der Kunstakademie. "
        "Das Residenzprogramm richtet sich an profilierte Druckgrafikerinnen und -grafiker "
        "und sieht die Nutzung der Druckwerkstatt im Hauptgebäude Gerberau für einen Zeitraum "
        "von drei bis sechs Monaten vor. Prof. Anita Gruber (KL-01) übernahm die künstlerische "
        "Leitung der Druckgrafik-Residenz und betreute das begleitende Kursprogramm (KU-003).",
        BODY))
    story.append(Paragraph(
        "Die Druckgrafik-Residenz wurde von der Stiftung Kunstförderung Oberrhein mit "
        "projektgebundenen Mitteln unterstützt (EV-01 / EV-07). Das Residenzprogramm "
        "gilt als besonderer programmatischer Erfolg des Berichtsjahres.",
        BODY))

    story.append(Paragraph("2.2 Sommeratelier-Workshopwoche (EV-06)", H2))
    story.append(Paragraph(
        "Vom 17. bis 23. Juli 2025 fand an allen drei Standorten (GER, WIE, STT) die "
        "Sommeratelier-Workshopwoche statt (EV-06). Sechs Intensivworkshops wurden durchgeführt, "
        "betreut von KL-01 Prof. Gruber (Druckgrafik), KL-03 Ingeborg Zellner (Aquarell), "
        "KL-05 Susanne Wältermann (Skulptur) und KL-06 Florian Neugebauer (Digitale Medien). "
        "Die Auslastung des Außenateliers Wiehre erreichte einen Rekordwert. "
        "Insgesamt nahmen rund 60 Teilnehmerinnen und Teilnehmer teil.",
        BODY))

    story.append(Paragraph('2.3 Frühjahrsausstellung „Form und Farbe“ (EV-03)', H2))
    story.append(Paragraph(
        "Am 3. März 2025 eröffnete im Gerberau-Ausstellungssaal die Frühjahrsausstellung "
        '„Form und Farbe“ (EV-03). Anwesend waren Vereinsvorsitzende Renate Quast (PER-001), '
        "Ausstellungskoordinatorin Franziska Oppelt (PER-009), Prof. Anita Gruber (KL-01) "
        "sowie Roland Kleiber (KL-08) und rund 112 Gäste. "
        "Im Rahmen der Eröffnung wurde der Verwendungsnachweis im Werkförderverfahren "
        "des Landes Baden-Württemberg eingereicht. Das Werkförderverfahren erforderte "
        "die Vorlage vollständiger Projektdokumentationen sowie den Nachweis der "
        "Mittelverwendung gegenüber der Bewilligungsbehörde. Die Ausstellung lief bis 30. April 2025.",
        BODY))

    story.append(Paragraph('2.4 Ausstellungseröffnung „Lichtwege" (EV-09)', H2))
    story.append(Paragraph(
        'Am 5. November 2025 eröffnete im Gerberau-Saal die Ausstellung „Lichtwege" (EV-09). '
        "Franziska Oppelt (PER-009) koordinierte die Eröffnung, an der Pavlos Demetriou (KL-04), "
        "Marta Szymańska (KL-07) und 89 Gäste teilnahmen. Erstmals wurde im Zuge dieser "
        "Ausstellung ein Atelierausleihe-Protokoll für eine externe Künstlerin angelegt.",
        BODY))

    story.append(PageBreak())

    # ── Page 3: Finanzkurzübersicht ──────────────────────────────────────────
    story.append(Paragraph("3. Finanzkurzübersicht 2025", H1))
    story.append(Paragraph(
        "Das Jahresbudget 2025 beläuft sich auf 1.240.000 € (Einnahmen = Ausgaben). "
        "Als gemeinnützige Organisation ist die Rheintal Akademie gGmbH zur zeitnahen "
        "Mittelverwendung nach §55 Abs. 1 Nr. 5 AO verpflichtet. Das Jahresergebnis beträgt 0 €.",
        BODY))

    fin_data = [
        ["EINNAHMEN", "Betrag (€)", "Anteil (%)"],
        ["Kursgebühren", "644.800", "52,0 %"],
        ["Mitgliedsbeiträge", "99.200", "8,0 %"],
        ["Öffentliche Fördergelder (Land BW + Stadt Freiburg)", "384.400", "31,0 %"],
        ["Stiftungsmittel projektgebunden", "111.600", "9,0 %"],
        ["Summe Einnahmen", "1.240.000", "100,0 %"],
    ]

    fin_table = Table(fin_data, colWidths=[9.5 * cm, 3 * cm, 2.5 * cm])
    fin_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), NAVY_RL),
        ("TEXTCOLOR", (0, 0), (-1, 0), white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("ROWBACKGROUNDS", (0, 1), (-1, -2), [HexColor("#f5f5f5"), white]),
        ("BACKGROUND", (0, -1), (-1, -1), ORANGE_RL),
        ("TEXTCOLOR", (0, -1), (-1, -1), white),
        ("FONTNAME", (0, -1), (-1, -1), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#cccccc")),
        ("ALIGN", (1, 0), (-1, -1), "RIGHT"),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(fin_table)
    story.append(Spacer(1, 0.4 * cm))

    aus_data = [
        ["AUSGABEN", "Betrag (€)", "Anteil (%)"],
        ["Personalaufwand Festangestellte", "682.000", "55,0 %"],
        ["Honorare Kursleiter", "198.400", "16,0 %"],
        ["Raum / Betrieb (Miete, Nebenkosten, Instandhaltung)", "198.400", "16,0 %"],
        ["Programm / Material (Werkstoffe, Ausstellungskosten)", "99.200", "8,0 %"],
        ["Verwaltung / Öffentlichkeitsarbeit", "62.000", "5,0 %"],
        ["Summe Ausgaben", "1.240.000", "100,0 %"],
    ]

    aus_table = Table(aus_data, colWidths=[9.5 * cm, 3 * cm, 2.5 * cm])
    aus_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), NAVY_RL),
        ("TEXTCOLOR", (0, 0), (-1, 0), white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("ROWBACKGROUNDS", (0, 1), (-1, -2), [HexColor("#f5f5f5"), white]),
        ("BACKGROUND", (0, -1), (-1, -1), ORANGE_RL),
        ("TEXTCOLOR", (0, -1), (-1, -1), white),
        ("FONTNAME", (0, -1), (-1, -1), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#cccccc")),
        ("ALIGN", (1, 0), (-1, -1), "RIGHT"),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(aus_table)
    story.append(Spacer(1, 0.3 * cm))
    story.append(Paragraph(
        "Jahresergebnis: 0 € — Die Organisation ist gemeinnützig; ein Jahresüberschuss "
        "ist weder angestrebt noch zulässig.",
        sty("KARFinNote", fontSize=9, fontName="Helvetica-Bold", textColor=NAVY_RL)))

    story.append(PageBreak())

    # ── Page 4: Fördernachweis ───────────────────────────────────────────────
    story.append(Paragraph("4. Fördernachweis 2025", H1))
    story.append(Paragraph(
        "Die öffentliche Förderung des Landes Baden-Württemberg wurde für das Haushaltsjahr 2025 "
        "mit Bescheid vom 9. Juni 2025 bewilligt (EV-05, adressiert an Urte Hamann, PER-011). "
        "Die Zuwendung in Höhe von 228.000 € dient der Finanzierung des laufenden Betriebs "
        "sowie anteiliger Personalkosten. Der Verwendungsnachweis ist bis zum 31. Januar 2026 "
        "bei der Bewilligungsbehörde einzureichen.",
        BODY))
    story.append(Paragraph(
        "Das Werkförderverfahren des Landes Baden-Württemberg erfordert die Vorlage "
        "vollständiger Projektdokumentationen sowie den Nachweis der Mittelverwendung. "
        'Im Rahmen des Werkförderverfahrens wurde für die Ausstellung „Form und Farbe” (EV-03) '
        "ein Projektbericht eingereicht, unterzeichnet von PER-002 Feuerbach und PER-011 Hamann.",
        BODY))

    foerd_data = [
        ["Geldgeber", "Art", "Betrag (€)", "Verwendungszweck", "Nachweis-Frist"],
        ["Land Baden-Württemberg", "Institutionelle Förderung", "228.000",
         "Laufender Betrieb, Personalkosten", "31.01.2026"],
        ["Stadt Freiburg im Breisgau", "Projektförderung Kulturprogramm", "156.400",
         "Ausstellungen, Bildungsprojekte", "28.02.2026"],
        ["Stiftung Kunstförderung Oberrhein", "Projektgebundene Mittel", "68.000",
         "Druckgrafik-Residenz (EV-01 / EV-07)", "30.06.2026"],
        ["Deutsche Kulturstiftung Süd", "Projektgebundene Mittel", "43.600",
         "Fotografie-Reihe Außenatelier", "30.09.2026"],
        ["Summe Förderung", "", "496.000", "", ""],
    ]

    foerd_table = Table(
        foerd_data,
        colWidths=[4.2 * cm, 3.4 * cm, 2.0 * cm, 3.8 * cm, 1.8 * cm],
    )
    foerd_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), NAVY_RL),
        ("TEXTCOLOR", (0, 0), (-1, 0), white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 7.5),
        ("ROWBACKGROUNDS", (0, 1), (-1, -2), [HexColor("#f5f5f5"), white]),
        ("BACKGROUND", (0, -1), (-1, -1), ORANGE_RL),
        ("TEXTCOLOR", (0, -1), (-1, -1), white),
        ("FONTNAME", (0, -1), (-1, -1), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#cccccc")),
        ("ALIGN", (2, 0), (2, -1), "RIGHT"),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
        ("TOPPADDING", (0, 0), (-1, -1), 2),
    ]))
    story.append(foerd_table)
    story.append(Spacer(1, 0.4 * cm))
    story.append(Paragraph(
        "Jahreshauptversammlung: Am 22. Oktober 2025 (EV-08) genehmigte die Mitgliederversammlung "
        "mit 318 anwesenden oder vertretenen Mitgliedern den Jahresabschluss 2024.",
        BODY))
    story.append(Spacer(1, 0.3 * cm))
    story.append(Paragraph(
        "Freiburg im Breisgau, Dezember 2025 — Renate Quast (PER-001), Vereinsvorsitzende | "
        "Dr. Margit Feuerbach (PER-002), Geschäftsführerin gGmbH | "
        "Lukas Endres (PER-003), Programmdirektor",
        SMALL))

    # ── Canvas class: header + footer on every page ──────────────────────────
    class KARHFC(rl_canvas.Canvas):
        """Canvas subclass that draws the KAR header+footer on every page.

        Only showPage() is overridden so that headers/footers are drawn before
        each page is committed. save() is NOT overridden — overriding it would
        draw on a blank extra page that ReportLab opens after the last showPage().
        """
        def __init__(self, filename, **kwargs):
            super().__init__(filename, **kwargs)
            self._pn = 0

        def showPage(self):
            self._pn += 1
            _draw_kar_page_header(self, PAGE_W, PAGE_H, doc_title)
            _draw_kar_page_footer(self, PAGE_W, self._pn)
            super().showPage()

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=A4,
        leftMargin=MARGIN, rightMargin=MARGIN,
        topMargin=3.5 * cm, bottomMargin=2.5 * cm,
        title=doc_title,
        author="Kunstakademie Rheintal e. V.",
    )
    doc.build(story, canvasmaker=KARHFC)


# ---------------------------------------------------------------------------
# 2. honorarvertrag_gruber.docx
# ---------------------------------------------------------------------------

def build_honorarvertrag(out_path: Path) -> None:
    """Honorar-Rahmenvertrag zwischen Rheintal Akademie gGmbH (PER-002) und KL-01 Prof. Gruber.

    BM25 target term: Atelierüberlassung (≥ 3× in §4, explicitly defined).
    Reference event: EV-02 (14.02.2025).
    Structure: Title + Präambel + §1–§6 + Unterschriftenblock + Anhang Lehrplan.
    §2 contains a Markdown-style table of Stundensätze.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    # ── Title ──────────────────────────────────────────────────────────────
    title = doc.add_paragraph("Honorar-Rahmenvertrag — Lehrtätigkeit Druckgrafik", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph(
        "Rheintal Akademie gemeinnützige GmbH — Prof. Anita Gruber (KL-01) | "
        "Datum: 14.02.2025 (EV-02) | Ort: Freiburg im Breisgau"
    )
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.runs[0].italic = True
    meta.runs[0].font.size = Pt(10)
    meta.runs[0].font.color.rgb = RGBColor(33, 52, 82)

    # ── Präambel ───────────────────────────────────────────────────────────
    doc.add_heading("Präambel", level=1)
    doc.add_paragraph(
        "Die Rheintal Akademie gemeinnützige GmbH (nachfolgend „Akademie”), vertreten durch ihre "
        "Geschäftsführerin Dr. Margit Feuerbach (PER-002), Gerberau 12, 79098 Freiburg im Breisgau, "
        "und Frau Prof. Anita Gruber (KL-01), freischaffende Druckgrafikerin und Hochschullehrerin "
        "(nachfolgend „Kursleiterin”), schließen den nachfolgenden Honorar-Rahmenvertrag über die "
        "Lehrtätigkeit im Bereich Druckgrafik sowie über die damit verbundene Nutzung der "
        "Infrastruktur der Akademie. Dieser Vertrag tritt am 14. Februar 2025 (EV-02) in Kraft "
        "und ersetzt alle vorherigen mündlichen oder schriftlichen Absprachen zwischen den Parteien."
    )

    # ── §1 Vertragsgegenstand ──────────────────────────────────────────────
    doc.add_heading("§1 Vertragsgegenstand", level=2)
    doc.add_paragraph(
        "Gegenstand dieses Vertrages ist die selbstständige Lehrtätigkeit der Kursleiterin "
        "im Fachbereich Druckgrafik (Kürzel DRU) an der Kunstakademie Rheintal e. V. und der "
        "Rheintal Akademie gGmbH. Die Kursleiterin übernimmt die eigenverantwortliche Planung, "
        "Vorbereitung und Durchführung der ihr zugewiesenen Kurse gemäß dem jeweils gültigen "
        "Trimesterprogramm der Akademie."
    )
    doc.add_paragraph(
        "Die Kursleiterin ist im Rahmen dieses Vertrages nicht als Arbeitnehmerin tätig. "
        "Sie ist in der Gestaltung ihrer Unterrichtsmethoden frei, soweit sie den inhaltlichen "
        "Rahmen der Kursausschreibung einhält. Die Akademie stellt die Kursräume und die "
        "erforderliche Grundausstattung zur Verfügung; Spezialmaterial kann nach Absprache "
        "aus dem Akademiebestand entnommen oder auf Kosten der Akademie beschafft werden."
    )

    # ── §2 Vergütung ──────────────────────────────────────────────────────
    doc.add_heading("§2 Vergütung", level=2)
    doc.add_paragraph(
        "Die Kursleiterin erhält für ihre Lehrtätigkeit ein Honorar gemäß der nachfolgenden "
        "Stundensatztabelle. Grundlage der Abrechnung ist die tatsächlich geleistete Unterrichtszeit "
        "inklusive unmittelbarer Vor- und Nachbereitung (max. 15 % Pauschale). Die Abrechnung "
        "erfolgt monatlich auf Basis einer von der Kursleiterin eingereichten Honorarnote."
    )

    # Stundensatz table
    rate_table = doc.add_table(rows=5, cols=3)
    rate_table.style = "Table Grid"
    headers = ["Kurstyp", "Stundensatz (€)", "Anmerkung"]
    hdr_row = rate_table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "213452")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:val"), "clear")
        tc_pr.append(shd)

    rate_rows = [
        ("Einführungskurs Druckgrafik", "85 €", "Standardsatz gemäß Honorarordnung"),
        ("Vertiefungskurs Druckgrafik", "90 €", "Erhöhter Satz bei ≥ 8 Std./Kurs"),
        ("Intensivworkshop (Sommeratelier / Residenz)", "95 €", "Residenz- und Workshopformate"),
        ("Gastvorlesung / Einzelveranstaltung", "85 €", "Gleichgestellt mit Einführungskurs"),
    ]
    for r_idx, (kurs, satz, anm) in enumerate(rate_rows):
        row = rate_table.rows[r_idx + 1]
        row.cells[0].text = kurs
        row.cells[1].text = satz
        row.cells[2].text = anm

    doc.add_paragraph("")

    doc.add_paragraph(
        "Reisekosten werden nur erstattet, soweit sie vorab schriftlich durch die Geschäftsführerin "
        "Dr. Margit Feuerbach (PER-002) genehmigt wurden. Alle Honorare sind zuzüglich der "
        "gesetzlichen Umsatzsteuer, sofern die Kursleiterin zur Umsatzsteuer optiert."
    )

    # ── §3 Laufzeit ──────────────────────────────────────────────────────
    doc.add_heading("§3 Laufzeit", level=2)
    doc.add_paragraph(
        "Dieser Vertrag gilt ab dem 14. Februar 2025 (EV-02) und hat eine Grundlaufzeit bis zum "
        "31. Dezember 2025. Er verlängert sich automatisch um jeweils ein Kalenderjahr, sofern "
        "keine der Parteien ihn bis spätestens drei Monate vor Ablauf schriftlich kündigt."
    )
    doc.add_paragraph(
        "Für das Trimester 3 (September bis Dezember 2025) wurde das Residenzstipendium der "
        "Kursleiterin im Rahmen des T3-Kickoffs (EV-07, 12.09.2025) verlängert. "
        "Die Kursleiterin betreut im laufenden Programmjahr 2025 vier aktive Kurse "
        "(KU-001, KU-002, KU-003 sowie einen weiteren Druckgrafikkurs)."
    )
    doc.add_paragraph(
        "Anpassungen der Honorarsätze werden im gegenseitigen Einvernehmen schriftlich "
        "vereinbart und gelten ab dem nächsten vollen Trimesterbeginn. Bereits laufende "
        "Kurse werden zu den bei Kursbeginn vereinbarten Sätzen abgerechnet."
    )

    # ── §4 Atelierüberlassung ─────────────────────────────────────────────
    doc.add_heading("§4 Atelierüberlassung", level=2)
    doc.add_paragraph(
        "Die Atelierüberlassung regelt das Recht der Kursleiterin, die Druckwerkstatt im "
        "Hauptgebäude Gerberau (GER) außerhalb der regulären Kurszeiten für ihre eigene "
        "künstlerische Arbeit unentgeltlich zu nutzen. Die Atelierüberlassung ist auf "
        "maximal 8 Stunden pro Woche begrenzt und muss mindestens 48 Stunden im Voraus "
        "schriftlich bei Gundula Heth (PER-007, Atelierleitung Gerberau) angemeldet werden."
    )
    doc.add_paragraph(
        "Die Atelierüberlassung umfasst die Nutzung aller fest installierten Druckpressen, "
        "Walzen und Säurewannen der Druckwerkstatt. Verbrauchsmaterialien (Farben, Papier, "
        "Ätzchemikalien) sind von der Kursleiterin auf eigene Kosten zu beschaffen oder "
        "gegen ein pauschales Materialentgelt von 15 € pro Nutzungstag aus dem Akademiebestand "
        "zu entnehmen. Die Kursleiterin hat die Pflicht, den Arbeitsbereich nach jeder Nutzung "
        "ordnungsgemäß zu reinigen und den Zustand gemäß dem Übergabeprotokoll zu dokumentieren."
    )
    doc.add_paragraph(
        "Im Falle einer Verletzung der Pflichten aus der Atelierüberlassung — insbesondere "
        "bei wiederholter Überschreitung der Wochenstundengrenze oder schuldhafter Beschädigung "
        "von Geräten — kann die Akademie die Atelierüberlassung mit sofortiger Wirkung aussetzen. "
        "Eine Wiederherstellung des Rechts auf Atelierüberlassung ist erst nach schriftlicher "
        "Abstimmung mit der Geschäftsführerin Dr. Margit Feuerbach (PER-002) möglich."
    )
    doc.add_paragraph(
        "Werke, die im Rahmen der Atelierüberlassung entstehen, sind ausschließlich Eigentum "
        "der Kursleiterin. Die Akademie erwirbt durch die Atelierüberlassung keinerlei Nutzungs-, "
        "Verwertungs- oder Ausstellungsrechte an den entstehenden Werken, sofern nicht gesondert "
        "schriftlich vereinbart."
    )

    # ── §5 Urheberrecht und Verwertung ────────────────────────────────────
    doc.add_heading("§5 Urheberrecht und Verwertung", level=2)
    doc.add_paragraph(
        "Sämtliche im Rahmen der Lehrtätigkeit erstellten Kursmaterialien (Handouts, "
        "Arbeitsblätter, digitale Präsentationen) verbleiben im Eigentum der Kursleiterin. "
        "Die Akademie erhält ein nicht-exklusives, nicht übertragbares Nutzungsrecht "
        "für die interne Verwendung in den vereinbarten Kursen."
    )
    doc.add_paragraph(
        "Die Kursleiterin stimmt zu, dass die Akademie Fotos und kurze Videoaufnahmen "
        "des Kursbetriebs für Öffentlichkeitsarbeit und Programmwerbung verwenden darf, "
        "sofern die Kursleiterin persönlich abgebildet ist und der Verwendung im Einzelfall "
        "zugestimmt hat. Eine kommerzielle Verwertung außerhalb der Akademiekommunikation "
        "ist ausgeschlossen."
    )

    # ── §6 Kündigung und Schlussbestimmungen ──────────────────────────────
    doc.add_heading("§6 Kündigung und Schlussbestimmungen", level=2)
    doc.add_paragraph(
        "Dieser Vertrag kann von beiden Parteien mit einer Frist von drei Monaten zum "
        "Trimesterende schriftlich gekündigt werden. Eine außerordentliche Kündigung aus "
        "wichtigem Grund ist jederzeit möglich. Als wichtiger Grund gilt insbesondere die "
        "wiederholte und schuldhafte Verletzung vertraglicher Pflichten oder die dauerhaft "
        "mangelnde Eignung für die vereinbarten Lehrleistungen."
    )
    doc.add_paragraph(
        "Änderungen und Ergänzungen dieses Vertrages bedürfen der Schriftform. "
        "Mündliche Nebenabreden bestehen nicht. Sollte eine Bestimmung dieses Vertrages "
        "unwirksam sein, bleibt der übrige Vertrag davon unberührt; die unwirksame Bestimmung "
        "ist durch eine wirksame zu ersetzen, die dem wirtschaftlichen Zweck der ursprünglichen "
        "Regelung möglichst nahekommt."
    )
    doc.add_paragraph(
        "Erfüllungsort und Gerichtsstand ist Freiburg im Breisgau. Es gilt das Recht der "
        "Bundesrepublik Deutschland. Beide Parteien erhalten ein Original dieses Vertrages."
    )

    # ── Unterschriftenblock ───────────────────────────────────────────────
    doc.add_paragraph("")
    doc.add_heading("Unterschriften", level=1)

    sign_data = [
        ["Datum:", "14.02.2025 (EV-02)"],
        ["Ort:", "Freiburg im Breisgau"],
        ["Für die Rheintal Akademie gGmbH:", "Dr. Margit Feuerbach (PER-002), Geschäftsführerin"],
        ["Kursleiterin:", "Prof. Anita Gruber (KL-01)"],
    ]
    sign_table = doc.add_table(rows=4, cols=2)
    sign_table.style = "Table Grid"
    for r_idx, (label, value) in enumerate(sign_data):
        row = sign_table.rows[r_idx]
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].text = value

    doc.add_paragraph("")
    doc.add_paragraph(
        "_________________________                    _________________________"
    )
    doc.add_paragraph(
        "Dr. Margit Feuerbach (PER-002)               Prof. Anita Gruber (KL-01)"
    ).runs[0].italic = True

    # ── Anhang: Lehrplan ──────────────────────────────────────────────────
    doc.add_paragraph("")
    doc.add_heading("Anhang: Lehrplan Druckgrafik-Kurse", level=1)
    doc.add_paragraph(
        "Nachfolgend sind die Semesterinhalte aufgeführt, die Prof. Anita Gruber (KL-01) "
        "in den ihr zugewiesenen Kursen zu vermitteln hat. Die Inhalte sind als Rahmenplan "
        "zu verstehen; die methodische Umsetzung liegt im Ermessen der Kursleiterin."
    )

    lehrplan_items = [
        "Einführung in die Geschichte und Technik der Druckgrafik: Hochdruck, Tiefdruck, Flachdruck, Siebdruck",
        "Radierung: Vorbereitung der Druckplatte, Ätzverfahren, Drucktechnik, Reinigung und Lagerung",
        "Holzschnitt und Linolschnitt: Werkzeugkunde, Schnitttechniken, Farbregisterdruck",
        "Experimentelle Druckgrafik: Mixed-Media-Ansätze, Monotypie, Transferdrucktechniken",
        "Druckwerkstatt-Management: Sicherheitsunterweisung, Gerätepflege, Umgang mit Druckchemikalien",
        "Werkpräsentation und Portfolioaufbau: Dokumentation eigener Arbeiten, Ausstellungsvorbereitung",
        "Einführung in digitale Druckverfahren und deren Kombination mit analogen Techniken",
    ]

    for item in lehrplan_items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


# ---------------------------------------------------------------------------
# 3. haushaltsplan_2025.xlsx
# ---------------------------------------------------------------------------

def build_haushaltsplan(out_path: Path) -> None:
    """3-sheet Haushaltsplan 2025 der Kunstakademie Rheintal e. V. / Rheintal Akademie gGmbH.

    Sheet 1: Einnahmen-Ausgaben (incl. Satzungsklausel-7)
    Sheet 2: Förderübersicht (incl. Verwendungsnachweis-Frist ≥ 2×)
    Sheet 3: Honorarübersicht (34 KL rows, Summe = 198.400 €)
    BM25 target terms: Satzungsklausel-7 (≥ 1×), Verwendungsnachweis-Frist (≥ 2×).
    References: PER-002 Feuerbach, PER-005 Schipper, PER-011 Hamann.
    """
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    wb = Workbook()

    HEADER_FILL = PatternFill("solid", fgColor="213452")
    HEADER_FONT = Font(bold=True, color="FFFFFF", name="Calibri")
    BOLD_FONT = Font(bold=True, name="Calibri")
    NORMAL_FONT = Font(name="Calibri")
    TOTAL_FILL = PatternFill("solid", fgColor="EE7F00")
    TOTAL_FONT = Font(bold=True, color="FFFFFF", name="Calibri")
    EUR_FMT = '#,##0 "€"'

    double_bottom = Border(
        bottom=Side(style="double"),
        top=Side(style="thin"),
    )

    def col_width(ws, col, width):
        ws.column_dimensions[get_column_letter(col)].width = width

    # ── Sheet 1: Einnahmen-Ausgaben ──────────────────────────────────────
    ws1 = wb.active
    ws1.title = "Einnahmen-Ausgaben"

    ws1["A1"] = ("Kunstakademie Rheintal e. V. / Rheintal Akademie gGmbH "
                 "— Haushaltsplan 2025")
    ws1["A1"].font = Font(bold=True, size=13, name="Calibri")
    ws1["A2"] = "Verantwortlich: PER-005 Tobias Schipper (Finanzbuchhaltung)"
    ws1["A2"].font = Font(italic=True, name="Calibri")
    ws1["A3"] = "Geschäftsführerin gGmbH: PER-002 Dr. Margit Feuerbach | Drittmittel: PER-011 Urte Hamann"
    ws1["A3"].font = Font(italic=True, name="Calibri")

    # Einnahmen block
    einnahmen_header_row = 5
    ws1.cell(einnahmen_header_row, 1, "EINNAHMEN").font = Font(bold=True, size=11, name="Calibri")

    einnahmen_rows = [
        ("Position", "Betrag (€)"),
        ("Kursgebühren", 644_800),
        ("Mitgliedsbeiträge", 99_200),
        ("Öffentliche Fördergelder (Land BW + Stadt Freiburg)", 384_400),
        ("Stiftungsmittel projektgebunden", 111_600),
        ("Summe Einnahmen", 1_240_000),
    ]

    r = einnahmen_header_row + 1
    for i, (pos, val) in enumerate(einnahmen_rows):
        ws1.cell(r, 1, pos)
        ws1.cell(r, 2, val if isinstance(val, int) else val)
        if i == 0:
            ws1.cell(r, 1).font = HEADER_FONT
            ws1.cell(r, 1).fill = HEADER_FILL
            ws1.cell(r, 2).font = HEADER_FONT
            ws1.cell(r, 2).fill = HEADER_FILL
        elif pos.startswith("Summe"):
            ws1.cell(r, 1).font = TOTAL_FONT
            ws1.cell(r, 1).fill = TOTAL_FILL
            ws1.cell(r, 2).font = TOTAL_FONT
            ws1.cell(r, 2).fill = TOTAL_FILL
            ws1.cell(r, 2).number_format = EUR_FMT
            ws1.cell(r, 1).border = double_bottom
            ws1.cell(r, 2).border = double_bottom
        else:
            ws1.cell(r, 2).number_format = EUR_FMT
        r += 1

    # Ausgaben block (one blank row gap)
    ausgaben_header_row = r + 1
    ws1.cell(ausgaben_header_row, 1, "AUSGABEN").font = Font(bold=True, size=11, name="Calibri")

    ausgaben_rows = [
        ("Position", "Betrag (€)"),
        ("Personalaufwand Festangestellte", 682_000),
        ("Honorare Kursleiter", 198_400),
        ("Raum / Betrieb (Miete, Nebenkosten, Instandhaltung)", 198_400),
        ("Programm / Material (Werkstoffe, Ausstellungskosten)", 99_200),
        ("Verwaltung / Öffentlichkeitsarbeit", 62_000),
        ("Summe Ausgaben", 1_240_000),
    ]

    r = ausgaben_header_row + 1
    for i, (pos, val) in enumerate(ausgaben_rows):
        ws1.cell(r, 1, pos)
        ws1.cell(r, 2, val if isinstance(val, int) else val)
        if i == 0:
            ws1.cell(r, 1).font = HEADER_FONT
            ws1.cell(r, 1).fill = HEADER_FILL
            ws1.cell(r, 2).font = HEADER_FONT
            ws1.cell(r, 2).fill = HEADER_FILL
        elif pos.startswith("Summe"):
            ws1.cell(r, 1).font = TOTAL_FONT
            ws1.cell(r, 1).fill = TOTAL_FILL
            ws1.cell(r, 2).font = TOTAL_FONT
            ws1.cell(r, 2).fill = TOTAL_FILL
            ws1.cell(r, 2).number_format = EUR_FMT
            ws1.cell(r, 1).border = double_bottom
            ws1.cell(r, 2).border = double_bottom
        else:
            ws1.cell(r, 2).number_format = EUR_FMT
        r += 1

    # Jahresergebnis row
    r += 1
    ws1.cell(r, 1, "Jahresergebnis").font = Font(bold=True, name="Calibri")
    ws1.cell(r, 2, 0).number_format = EUR_FMT
    ws1.cell(r, 2).font = Font(bold=True, name="Calibri")

    r += 1
    ws1.cell(r, 1, (
        "Hinweis: Rücklage gemäß Satzungsklausel-7 wird aus Programm/Material-Überschüssen "
        "gebildet (Beschluss EV-10, 10.12.2025). Jahresergebnis 0 € — gemeinnützig."
    ))
    ws1.cell(r, 1).font = Font(italic=True, color="CC0000", name="Calibri")

    col_width(ws1, 1, 60)
    col_width(ws1, 2, 18)

    # ── Sheet 2: Förderübersicht ──────────────────────────────────────────
    ws2 = wb.create_sheet("Förderübersicht")

    ws2["A1"] = "Kunstakademie Rheintal e. V. — Förderübersicht 2025"
    ws2["A1"].font = Font(bold=True, size=13, name="Calibri")
    ws2["A2"] = (
        "Verantwortlich: PER-011 Urte Hamann (Förderanträge / Drittmittel) | "
        "Verwendungsnachweis-Frist je Geldgeber beachten."
    )
    ws2["A2"].font = Font(italic=True, name="Calibri")

    foerd_headers = ["Geldgeber", "Art", "Betrag (€)", "Verwendungszweck",
                     "Verwendungsnachweis-Frist", "Status"]
    foerd_data = [
        ("Land Baden-Württemberg", "Institutionelle Förderung", 228_000,
         "Laufender Betrieb, Personalkosten", "31.01.2026", "Bescheid EV-05"),
        ("Stadt Freiburg im Breisgau", "Projektförderung Kulturprogramm", 96_400,
         "Ausstellungen, Bildungsprojekte", "30.06.2026", "Bewilligt"),
        ("Stiftung Kulturraum Süd", "Projektgebundene Mittel", 60_000,
         "Druckgrafik-Residenz (EV-01 / EV-07)", "30.06.2026", "Projektgebunden"),
        ("Stiftung Bildungsbrücke", "Projektgebundene Mittel", 51_600,
         "Fotografie-Reihe Außenatelier", "30.09.2026", "Projektgebunden"),
    ]

    header_row = ws2[4]
    for col_idx, h in enumerate(foerd_headers, start=1):
        cell = ws2.cell(4, col_idx, h)
        cell.font = HEADER_FONT
        cell.fill = HEADER_FILL
        cell.alignment = Alignment(horizontal="center")

    for r_idx, row_data in enumerate(foerd_data, start=5):
        geldgeber, art, betrag, zweck, frist, status = row_data
        ws2.cell(r_idx, 1, geldgeber)
        ws2.cell(r_idx, 2, art)
        ws2.cell(r_idx, 3, betrag).number_format = EUR_FMT
        ws2.cell(r_idx, 4, zweck)
        ws2.cell(r_idx, 5, frist)
        ws2.cell(r_idx, 6, status)
        for c in range(1, 7):
            ws2.cell(r_idx, c).font = NORMAL_FONT

    # Summe row
    summe_r = 5 + len(foerd_data)
    ws2.cell(summe_r, 1, "Summe Förderung").font = TOTAL_FONT
    ws2.cell(summe_r, 1).fill = TOTAL_FILL
    ws2.cell(summe_r, 3, 436_000).number_format = EUR_FMT
    ws2.cell(summe_r, 3).font = TOTAL_FONT
    ws2.cell(summe_r, 3).fill = TOTAL_FILL

    # Note about Verwendungsnachweis-Frist
    note_r = summe_r + 2
    ws2.cell(note_r, 1, (
        "Hinweis: Die Verwendungsnachweis-Frist ist für jeden Fördergeber verbindlich. "
        "Verspätete Einreichungen führen zu Rückforderungen. Koordination: PER-011 Urte Hamann."
    ))
    ws2.cell(note_r, 1).font = Font(italic=True, name="Calibri")

    col_width(ws2, 1, 32)
    col_width(ws2, 2, 28)
    col_width(ws2, 3, 16)
    col_width(ws2, 4, 36)
    col_width(ws2, 5, 24)
    col_width(ws2, 6, 20)

    # ── Sheet 3: Honorarübersicht ─────────────────────────────────────────
    ws3 = wb.create_sheet("Honorarübersicht")

    ws3["A1"] = "Honorarübersicht 2025 — Kursleiter KL-01 … KL-34"
    ws3["A1"].font = Font(bold=True, size=13, name="Calibri")
    ws3["A2"] = "Verantwortlich: PER-005 Tobias Schipper | Summe Gesamthonorar: 198.400 €"
    ws3["A2"].font = Font(italic=True, name="Calibri")

    hon_headers = ["KursleiterID", "Name", "Fachbereich",
                   "Anzahl_Kurse_2025", "Stunden_2025",
                   "Honorar_Stundensatz", "Gesamthonorar_2025"]
    for col_idx, h in enumerate(hon_headers, start=1):
        cell = ws3.cell(4, col_idx, h)
        cell.font = HEADER_FONT
        cell.fill = HEADER_FILL
        cell.alignment = Alignment(horizontal="center")

    # All 34 KL rows from RHEINTAL.md section 4 — Gesamthonorar sums to exactly 198.400 €.
    # Hours per KL: 4-course KL = 120 h (KL-01), 3-course KL = 120 h,
    # 2-course KL = 80 h, 1-course KL = 40 h.
    # KL-01 carries 120 h (not 160) to balance the total precisely.
    kl_rows = [
        # fmt: (ID, Name, Fachbereich, Anzahl_Kurse, Stunden, Satz, Gesamt)
        ("KL-01", "Prof. Anita Gruber",        "Druckgrafik",     4, 120, 85, 10200),
        ("KL-02", "Berthold Schantz",           "Ölmalerei",       3, 120, 75,  9000),
        ("KL-03", "Ingeborg Zellner",           "Aquarell",        3, 120, 70,  8400),
        ("KL-04", "Pavlos Demetriou",           "Fotografie",      3, 120, 80,  9600),
        ("KL-05", "Susanne Wältermann",         "Skulptur",        2,  80, 90,  7200),
        ("KL-06", "Florian Neugebauer",         "Digitale Medien", 3, 120, 80,  9600),
        ("KL-07", "Marta Szymańska",            "Mixed Media",     3, 120, 65,  7800),
        ("KL-08", "Roland Kleiber",             "Ölmalerei",       2,  80, 75,  6000),
        ("KL-09", "Yuki Tanigawa",              "Aquarell",        2,  80, 60,  4800),
        ("KL-10", "Dominic Faure",              "Kunstgeschichte", 2,  80, 70,  5600),
        ("KL-11", "Claudia Rennert",            "Aktzeichnen",     2,  80, 65,  5200),
        ("KL-12", "Sebastian Borck",            "Druckgrafik",     2,  80, 80,  6400),
        ("KL-13", "Margarete Füsslin",          "Keramik",         3, 120, 70,  8400),
        ("KL-14", "Hanno Winkelmann",           "Musikproduktion", 2,  80, 85,  6800),
        ("KL-15", "Antje Steudel",              "Skulptur",        2,  80, 90,  7200),
        ("KL-16", "Georgios Papadimitriou",     "Fotografie",      2,  80, 75,  6000),
        ("KL-17", "Nina Schwarzbach",           "Mixed Media",     2,  80, 65,  5200),
        ("KL-18", "Tanja Urbach",               "Ölmalerei",       2,  80, 70,  5600),
        ("KL-19", "Dieter Fleckenstein",        "Aktzeichnen",     2,  80, 60,  4800),
        ("KL-20", "Maria Kovács",               "Aquarell",        2,  80, 65,  5200),
        ("KL-21", "Andreas Langwald",           "Kunstgeschichte", 2,  80, 70,  5600),
        ("KL-22", "Veronika Söllner",           "Keramik",         2,  80, 75,  6000),
        ("KL-23", "Raphael Drouet",             "Druckgrafik",     2,  80, 80,  6400),
        ("KL-24", "Brigitte Unterweger",        "Mixed Media",     2,  80, 65,  5200),
        ("KL-25", "Thomas Hirtreiter",          "Digitale Medien", 2,  80, 80,  6400),
        ("KL-26", "Leila Moussavi",             "Musikproduktion", 2,  80, 85,  6800),
        ("KL-27", "Jan Overwien",               "Skulptur",        1,  40, 90,  3600),
        ("KL-28", "Silvia Pfanner",             "Ölmalerei",       1,  40, 70,  2800),
        ("KL-29", "Konrad Autenrieth",          "Aquarell",        1,  40, 60,  2400),
        ("KL-30", "Fatima Boussouf",            "Fotografie",      1,  40, 75,  3000),
        ("KL-31", "Erika Lindström",            "Kunstgeschichte", 1,  40, 70,  2800),
        ("KL-32", "Cyril Magnin",               "Aktzeichnen",     1,  40, 65,  2600),
        ("KL-33", "Özlem Demir",                "Keramik",         1,  40, 70,  2800),
        ("KL-34", "Patrick Sauerwein",          "Digitale Medien", 1,  40, 75,  3000),
    ]

    for r_idx, row_data in enumerate(kl_rows, start=5):
        kl_id, name, fach, kurse, stunden, satz, gesamt = row_data
        ws3.cell(r_idx, 1, kl_id)
        ws3.cell(r_idx, 2, name)
        ws3.cell(r_idx, 3, fach)
        ws3.cell(r_idx, 4, kurse)
        ws3.cell(r_idx, 5, stunden)
        ws3.cell(r_idx, 6, satz).number_format = EUR_FMT
        ws3.cell(r_idx, 7, gesamt).number_format = EUR_FMT

    # Summe row
    summe_kl_r = 5 + len(kl_rows)
    ws3.cell(summe_kl_r, 1, "Summe").font = TOTAL_FONT
    ws3.cell(summe_kl_r, 1).fill = TOTAL_FILL
    ws3.cell(summe_kl_r, 7, 198_400).number_format = EUR_FMT
    ws3.cell(summe_kl_r, 7).font = TOTAL_FONT
    ws3.cell(summe_kl_r, 7).fill = TOTAL_FILL

    col_width(ws3, 1, 12)
    col_width(ws3, 2, 28)
    col_width(ws3, 3, 20)
    col_width(ws3, 4, 18)
    col_width(ws3, 5, 14)
    col_width(ws3, 6, 22)
    col_width(ws3, 7, 22)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(out_path))


# ---------------------------------------------------------------------------
# 4. kursbelegung_legacy.xls
# ---------------------------------------------------------------------------

def build_kursbelegung_legacy(out_path: Path) -> None:
    """Legacy-Export Kursbelegungsdaten 2018–2022 im BIFF8 .xls Format via xlwt.

    Single sheet 'Kursbelegung_alt'. Flat raw export style — minimal styling.
    BM25 target term: Trimesterauslastungsindex (appears in a hint row at the top).
    Covers KL-01 (Druckgrafik), KL-03 (Keramik), KL-04 (Schwarz-Weiß-Fotografie).
    Trailing comment row references Buchungssystem v2.4 (Stand 31.12.2022).
    """
    import xlwt

    wb = xlwt.Workbook(encoding="utf-8")
    ws = wb.add_sheet("Kursbelegung_alt")

    bold = xlwt.easyxf("font: bold on")
    normal = xlwt.easyxf()

    ws.col(0).width = 5500
    ws.col(1).width = 9000
    ws.col(2).width = 4000
    ws.col(3).width = 7000
    ws.col(4).width = 4000
    ws.col(5).width = 3000
    ws.col(6).width = 5000
    ws.col(7).width = 5000
    ws.col(8).width = 7000

    # Hint row containing Trimesterauslastungsindex (BM25 target term)
    ws.write(0, 0, (
        "Hinweis: Trimesterauslastungsindex je Kurs als (Belegt/Max)*100 — "
        "Export Buchungssystem v2.4, Stand 31.12.2022"
    ), normal)

    # Column header row
    headers = [
        "KURS_ID", "KURSNAME", "LEITER_ID", "LEITER_NAME",
        "TRIMESTER", "JAHR", "TEILNEHMER_BELEGT", "TEILNEHMER_MAX",
        "AUSLASTUNG_PROZENT",
    ]
    for col_idx, h in enumerate(headers):
        ws.write(1, col_idx, h, bold)

    # Historical data 2018–2022
    # KL-01 = Druckgrafik (Prof. Gruber), KL-03 = Keramik (Ingeborg Zellner),
    # KL-04 = Schwarz-Weiß-Fotografie (Pavlos Demetriou in legacy)
    legacy_data = [
        # 2018
        ("LEG-2018-T1-001", "Radierung Grundkurs",          "KL-01", "A.Gruber",    "T1", 2018, 10, 12, 83),
        ("LEG-2018-T1-002", "Keramik Einführung",           "KL-03", "I.Zellner",   "T1", 2018,  8, 10, 80),
        ("LEG-2018-T1-003", "SW-Fotografie Grundkurs",      "KL-04", "P.Demetriou", "T1", 2018,  9, 10, 90),
        ("LEG-2018-T2-001", "Hochdruck Linolschnitt",       "KL-01", "A.Gruber",    "T2", 2018,  7, 10, 70),
        ("LEG-2018-T2-002", "Töpfern für Anfänger",         "KL-03", "I.Zellner",   "T2", 2018, 10, 12, 83),
        ("LEG-2018-T3-001", "Radierung Vertiefung",         "KL-01", "A.Gruber",    "T3", 2018,  8, 10, 80),
        ("LEG-2018-T3-002", "SW-Fotografie Dunkelkammer",   "KL-04", "P.Demetriou", "T3", 2018,  6,  8, 75),
        # 2019
        ("LEG-2019-T1-001", "Radierung Grundkurs",          "KL-01", "A.Gruber",    "T1", 2019, 11, 12, 92),
        ("LEG-2019-T1-002", "Keramik Einführung",           "KL-03", "I.Zellner",   "T1", 2019,  9, 10, 90),
        ("LEG-2019-T2-001", "Experimentelle Druckgrafik",   "KL-01", "A.Gruber",    "T2", 2019,  6,  8, 75),
        ("LEG-2019-T2-002", "Raku-Brennen",                 "KL-03", "I.Zellner",   "T2", 2019,  8, 10, 80),
        ("LEG-2019-T2-003", "SW-Fotografie Grundkurs",      "KL-04", "P.Demetriou", "T2", 2019, 10, 10, 100),
        ("LEG-2019-T3-001", "Hochdruck Holzschnitt",        "KL-01", "A.Gruber",    "T3", 2019,  8, 10, 80),
        ("LEG-2019-T3-002", "Keramik Vertiefung",           "KL-03", "I.Zellner",   "T3", 2019,  7, 10, 70),
        # 2020
        ("LEG-2020-T1-001", "Radierung Grundkurs",          "KL-01", "A.Gruber",    "T1", 2020,  5, 12, 42),
        ("LEG-2020-T1-002", "Keramik Einführung",           "KL-03", "I.Zellner",   "T1", 2020,  4, 10, 40),
        ("LEG-2020-T2-001", "SW-Fotografie Online",         "KL-04", "P.Demetriou", "T2", 2020,  8, 10, 80),
        ("LEG-2020-T3-001", "Radierung Vertiefung",         "KL-01", "A.Gruber",    "T3", 2020,  9, 12, 75),
        ("LEG-2020-T3-002", "Keramik Herbst",               "KL-03", "I.Zellner",   "T3", 2020,  7, 10, 70),
        # 2021
        ("LEG-2021-T1-001", "Radierung Grundkurs",          "KL-01", "A.Gruber",    "T1", 2021, 12, 12, 100),
        ("LEG-2021-T1-002", "Keramik Einführung",           "KL-03", "I.Zellner",   "T1", 2021, 10, 10, 100),
        ("LEG-2021-T2-001", "Linolschnitt Kompakt",         "KL-01", "A.Gruber",    "T2", 2021,  8, 10, 80),
        ("LEG-2021-T2-002", "SW-Fotografie Porträt",        "KL-04", "P.Demetriou", "T2", 2021,  9, 10, 90),
        ("LEG-2021-T3-001", "Experimentelle Druckgrafik",   "KL-01", "A.Gruber",    "T3", 2021,  7,  8, 88),
        ("LEG-2021-T3-002", "Keramik Weihnachten",          "KL-03", "I.Zellner",   "T3", 2021, 10, 12, 83),
        # 2022
        ("LEG-2022-T1-001", "Radierung Grundkurs",          "KL-01", "A.Gruber",    "T1", 2022, 12, 12, 100),
        ("LEG-2022-T1-002", "Keramik Einführung",           "KL-03", "I.Zellner",   "T1", 2022,  9, 10, 90),
        ("LEG-2022-T2-001", "SW-Fotografie Grundkurs",      "KL-04", "P.Demetriou", "T2", 2022, 10, 10, 100),
        ("LEG-2022-T2-002", "Hochdruck Holzschnitt",        "KL-01", "A.Gruber",    "T2", 2022,  8, 10, 80),
        ("LEG-2022-T3-001", "Keramik Vertiefung",           "KL-03", "I.Zellner",   "T3", 2022,  9, 10, 90),
        ("LEG-2022-T3-002", "Radierung Intensiv",           "KL-01", "A.Gruber",    "T3", 2022, 10, 12, 83),
        ("LEG-2022-T3-003", "SW-Fotografie Dunkelkammer",   "KL-04", "P.Demetriou", "T3", 2022,  8,  8, 100),
    ]

    for r_idx, row_data in enumerate(legacy_data, start=2):
        (kurs_id, kursname, leiter_id, leiter_name,
         trimester, jahr, belegt, max_tn, auslastung) = row_data
        ws.write(r_idx, 0, kurs_id, normal)
        ws.write(r_idx, 1, kursname, normal)
        ws.write(r_idx, 2, leiter_id, normal)
        ws.write(r_idx, 3, leiter_name, normal)
        ws.write(r_idx, 4, trimester, normal)
        ws.write(r_idx, 5, jahr, normal)
        ws.write(r_idx, 6, belegt, normal)
        ws.write(r_idx, 7, max_tn, normal)
        ws.write(r_idx, 8, auslastung, normal)

    # Trailing comment row
    trailing_r = 2 + len(legacy_data)
    ws.write(trailing_r, 0, (
        "Export: Buchungssystem v2.4 (Stand 31.12.2022, vor Migration auf neues System)"
    ), normal)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(out_path))


# ===========================================================================
# main
# ===========================================================================

def main() -> None:
    outputs = [
        (CORPUS_ROOT / "musterbau" / "finanzen_2025.xlsx", build_finanzen),
        (CORPUS_ROOT / "musterbau" / "finanzen_legacy.xls", build_finanzen_legacy_xls),
        (CORPUS_ROOT / "musterbau" / "memo_strategieklausur.docx", build_memo),
        (CORPUS_ROOT / "musterbau" / "geschaeftsbericht_2025.pdf", build_geschaeftsbericht),
        (CORPUS_ROOT / "long" / "handbuch_lang.pdf", build_long_handbook),
        (CORPUS_ROOT / "dedup" / "sample.pdf", build_dedup_sample),
        (CORPUS_ROOT / "phash" / "logo_repeating.pdf", build_phash_logo_repeating),
        (CORPUS_ROOT / "multidoc_a" / "appendix.pdf", build_multidoc_appendix),
        # ── Rheintal corpus ──────────────────────────────────────────────────
        (CORPUS_ROOT / "rheintal" / "taetigkeitsbericht_2025.pdf", build_taetigkeitsbericht),
        (CORPUS_ROOT / "rheintal" / "honorarvertrag_gruber.docx", build_honorarvertrag),
        (CORPUS_ROOT / "rheintal" / "haushaltsplan_2025.xlsx", build_haushaltsplan),
        (CORPUS_ROOT / "rheintal" / "kursbelegung_legacy.xls", build_kursbelegung_legacy),
    ]

    for out_path, builder in outputs:
        out_path.parent.mkdir(parents=True, exist_ok=True)
        builder(out_path)
        size = out_path.stat().st_size
        print(f"wrote: {out_path}  ({size:,} bytes)")

    print(f"\nTotal files written: {len(outputs)}")

    # Sanity check
    failures = []
    for out_path, _ in outputs:
        if not out_path.exists():
            failures.append(f"MISSING: {out_path}")
        elif out_path.stat().st_size < 1024:
            failures.append(f"TOO SMALL (<1 KB): {out_path}")
    if failures:
        print("\nSANITY CHECK FAILED:")
        for f in failures:
            print(f"  {f}")
        raise SystemExit(1)
    else:
        print(f"Sanity check passed: all {len(outputs)} files exist and are > 1 KB.")


if __name__ == "__main__":
    main()
